"""
📄 Ferramentas de PDF - Criação e Edição
Assistente Ranny V3

Funcionalidades:
- Criar PDF a partir de texto
- Mesclar múltiplos PDFs
- Extrair páginas específicas
- Comprimir PDF
- Adicionar marca d'água
- Converter imagens em PDF
"""

import io
import logging
import re
from datetime import datetime
from typing import List, Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)

# Importa bibliotecas de PDF
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF não instalado - algumas funções de PDF desabilitadas")

try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    logger.warning("ReportLab não instalado - criação de PDF desabilitada")

# Importa bibliotecas do Office
try:
    # Tenta importar skelmis (fork melhorado com imagens flutuantes)
    try:
        from skelmis.docx import Document
        from skelmis.docx.shared import Inches, Pt, Cm
        from skelmis.docx.enum.text import WD_ALIGN_PARAGRAPH
        from skelmis.docx.enum.style import WD_STYLE_TYPE
        HAS_SKELMIS_DOCX = True
        HAS_DOCX = True
        logger.info("Usando skelmis-python-docx (fork melhorado)")
    except ImportError:
        # Fallback para python-docx original
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        HAS_SKELMIS_DOCX = False
        HAS_DOCX = True
        logger.info("Usando python-docx original")
except ImportError:
    HAS_SKELMIS_DOCX = False
    HAS_DOCX = False
    logger.warning("python-docx não instalado - criação de DOCX desabilitada")

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False
    logger.warning("openpyxl não instalado - criação de XLSX desabilitada")


def criar_pdf_texto(texto: str, titulo: str = None) -> Optional[bytes]:
    """Cria um PDF a partir de texto
    
    Args:
        texto: Conteúdo do PDF
        titulo: Título opcional do documento
    
    Returns:
        bytes do PDF ou None se falhar
    """
    if not HAS_REPORTLAB:
        logger.error("ReportLab não instalado")
        return None
    
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Estilo personalizado para título
        titulo_style = ParagraphStyle(
            'TituloCustom',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            textColor=HexColor('#2C3E50'),
            alignment=1  # Centralizado
        )
        
        # Estilo para corpo do texto
        corpo_style = ParagraphStyle(
            'CorpoCustom',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=12,
            textColor=HexColor('#333333')
        )
        
        # Estilo para rodapé
        rodape_style = ParagraphStyle(
            'RodapeCustom',
            parent=styles['Normal'],
            fontSize=8,
            textColor=HexColor('#888888'),
            alignment=1
        )
        
        elementos = []
        
        # Título
        if titulo:
            elementos.append(Paragraph(titulo, titulo_style))
            elementos.append(Spacer(1, 20))
        
        # Processa o texto - divide em parágrafos
        paragrafos = texto.split('\n')
        for p in paragrafos:
            p = p.strip()
            if p:
                # Escapa caracteres especiais do HTML
                p = p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elementos.append(Paragraph(p, corpo_style))
            else:
                elementos.append(Spacer(1, 10))
        
        # Rodapé com data
        elementos.append(Spacer(1, 30))
        data_atual = datetime.now().strftime('%d/%m/%Y às %H:%M')
        elementos.append(Paragraph(f"Gerado em {data_atual} • Assistente Ranny", rodape_style))
        
        # Gera o PDF
        doc.build(elementos)
        
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        logger.info(f"PDF criado com sucesso: {len(pdf_bytes)} bytes")
        return pdf_bytes
        
    except Exception as e:
        logger.error(f"Erro ao criar PDF: {e}")
        return None


def mesclar_pdfs(pdfs: List[bytes]) -> Optional[bytes]:
    """Mescla múltiplos PDFs em um único arquivo
    
    Args:
        pdfs: Lista de bytes de PDFs
    
    Returns:
        bytes do PDF mesclado ou None se falhar
    """
    if not HAS_PYMUPDF:
        logger.error("PyMuPDF não instalado")
        return None
    
    if len(pdfs) < 2:
        logger.error("Precisa de pelo menos 2 PDFs para mesclar")
        return None
    
    try:
        # Cria documento de saída
        pdf_saida = fitz.open()
        
        for i, pdf_bytes in enumerate(pdfs):
            try:
                pdf_entrada = fitz.open(stream=pdf_bytes, filetype="pdf")
                pdf_saida.insert_pdf(pdf_entrada)
                pdf_entrada.close()
                logger.info(f"PDF {i+1} adicionado ({len(pdf_bytes)} bytes)")
            except Exception as e:
                logger.error(f"Erro ao processar PDF {i+1}: {e}")
                continue
        
        if pdf_saida.page_count == 0:
            logger.error("Nenhuma página foi adicionada")
            return None
        
        # Salva em bytes
        output = io.BytesIO()
        pdf_saida.save(output)
        pdf_saida.close()
        
        result = output.getvalue()
        output.close()
        
        logger.info(f"PDFs mesclados: {len(result)} bytes, {pdf_saida.page_count} páginas")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao mesclar PDFs: {e}")
        return None


def extrair_paginas(pdf_bytes: bytes, paginas: List[int]) -> Optional[bytes]:
    """Extrai páginas específicas de um PDF
    
    Args:
        pdf_bytes: bytes do PDF original
        paginas: Lista de números de página (1-indexed)
    
    Returns:
        bytes do PDF com as páginas extraídas ou None se falhar
    """
    if not HAS_PYMUPDF:
        logger.error("PyMuPDF não instalado")
        return None
    
    try:
        pdf_entrada = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_paginas = pdf_entrada.page_count
        
        # Valida páginas
        paginas_validas = []
        for p in paginas:
            if 1 <= p <= total_paginas:
                paginas_validas.append(p - 1)  # Converte para 0-indexed
            else:
                logger.warning(f"Página {p} inválida (total: {total_paginas})")
        
        if not paginas_validas:
            logger.error("Nenhuma página válida especificada")
            pdf_entrada.close()
            return None
        
        # Cria novo PDF com as páginas selecionadas
        pdf_saida = fitz.open()
        
        for page_num in paginas_validas:
            pdf_saida.insert_pdf(pdf_entrada, from_page=page_num, to_page=page_num)
        
        pdf_entrada.close()
        
        # Salva em bytes
        output = io.BytesIO()
        pdf_saida.save(output)
        pdf_saida.close()
        
        result = output.getvalue()
        output.close()
        
        logger.info(f"Páginas extraídas: {paginas_validas}, {len(result)} bytes")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao extrair páginas: {e}")
        return None


def comprimir_pdf(pdf_bytes: bytes, qualidade: str = "media") -> Optional[Tuple[bytes, dict]]:
    """Comprime um PDF reduzindo seu tamanho
    
    Args:
        pdf_bytes: bytes do PDF original
        qualidade: 'baixa', 'media' ou 'alta'
    
    Returns:
        Tuple (bytes do PDF comprimido, dict com info) ou None se falhar
    """
    if not HAS_PYMUPDF:
        logger.error("PyMuPDF não instalado")
        return None
    
    try:
        tamanho_original = len(pdf_bytes)
        
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        # Configurações de compressão baseadas na qualidade
        if qualidade == "baixa":
            # Máxima compressão, menor qualidade
            garbage = 4
            deflate = True
            clean = True
        elif qualidade == "alta":
            # Mínima compressão, maior qualidade
            garbage = 1
            deflate = True
            clean = False
        else:  # media
            garbage = 3
            deflate = True
            clean = True
        
        # Salva com compressão
        output = io.BytesIO()
        pdf.save(
            output,
            garbage=garbage,
            deflate=deflate,
            clean=clean,
            linear=True  # Otimiza para web
        )
        pdf.close()
        
        result = output.getvalue()
        output.close()
        
        tamanho_final = len(result)
        reducao = ((tamanho_original - tamanho_final) / tamanho_original) * 100
        
        info = {
            'tamanho_original': tamanho_original,
            'tamanho_final': tamanho_final,
            'reducao_percentual': round(reducao, 1),
            'qualidade': qualidade
        }
        
        logger.info(f"PDF comprimido: {tamanho_original} -> {tamanho_final} bytes ({reducao:.1f}% redução)")
        return result, info
        
    except Exception as e:
        logger.error(f"Erro ao comprimir PDF: {e}")
        return None


def adicionar_marca_dagua(pdf_bytes: bytes, texto: str, opacidade: float = 0.3) -> Optional[bytes]:
    """Adiciona marca d'água em todas as páginas do PDF
    
    Args:
        pdf_bytes: bytes do PDF original
        texto: Texto da marca d'água
        opacidade: Opacidade da marca (0.0 a 1.0)
    
    Returns:
        bytes do PDF com marca d'água ou None se falhar
    """
    if not HAS_PYMUPDF:
        logger.error("PyMuPDF não instalado")
        return None
    
    try:
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        for page in pdf:
            # Obtém dimensões da página
            rect = page.rect
            
            # Configura a marca d'água
            # Posiciona no centro da página, rotacionada
            center_x = rect.width / 2
            center_y = rect.height / 2
            
            # Adiciona texto como marca d'água
            page.insert_text(
                (center_x - len(texto) * 5, center_y),  # Posição aproximada
                texto,
                fontsize=50,
                fontname="helv",
                color=(0.5, 0.5, 0.5),  # Cinza
                rotate=45,  # Rotação diagonal
                overlay=True
            )
        
        # Salva em bytes
        output = io.BytesIO()
        pdf.save(output)
        pdf.close()
        
        result = output.getvalue()
        output.close()
        
        logger.info(f"Marca d'água adicionada: '{texto}'")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao adicionar marca d'água: {e}")
        return None


def imagens_para_pdf(imagens: List[bytes]) -> Optional[bytes]:
    """Converte uma ou mais imagens em PDF
    
    Args:
        imagens: Lista de bytes de imagens (PNG, JPG, etc)
    
    Returns:
        bytes do PDF ou None se falhar
    """
    if not HAS_PYMUPDF:
        logger.error("PyMuPDF não instalado")
        return None
    
    try:
        pdf = fitz.open()
        
        for i, img_bytes in enumerate(imagens):
            try:
                # Abre a imagem
                img = fitz.open(stream=img_bytes, filetype="png")
                
                # Converte para PDF
                pdf_bytes_temp = img.convert_to_pdf()
                img.close()
                
                # Abre o PDF temporário e insere no documento principal
                img_pdf = fitz.open(stream=pdf_bytes_temp, filetype="pdf")
                pdf.insert_pdf(img_pdf)
                img_pdf.close()
                
                logger.info(f"Imagem {i+1} adicionada ao PDF")
                
            except Exception as e:
                logger.error(f"Erro ao processar imagem {i+1}: {e}")
                continue
        
        if pdf.page_count == 0:
            logger.error("Nenhuma imagem foi convertida")
            return None
        
        # Salva em bytes
        output = io.BytesIO()
        pdf.save(output)
        pdf.close()
        
        result = output.getvalue()
        output.close()
        
        logger.info(f"Imagens convertidas para PDF: {len(result)} bytes, {pdf.page_count} páginas")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao converter imagens para PDF: {e}")
        return None


def get_pdf_info(pdf_bytes: bytes) -> Optional[dict]:
    """Obtém informações sobre um PDF
    
    Args:
        pdf_bytes: bytes do PDF
    
    Returns:
        dict com informações ou None se falhar
    """
    if not HAS_PYMUPDF:
        logger.error("PyMuPDF não instalado")
        return None
    
    try:
        pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        info = {
            'paginas': pdf.page_count,
            'tamanho_bytes': len(pdf_bytes),
            'tamanho_formatado': _formatar_tamanho(len(pdf_bytes)),
            'titulo': pdf.metadata.get('title', ''),
            'autor': pdf.metadata.get('author', ''),
            'criado': pdf.metadata.get('creationDate', ''),
            'modificado': pdf.metadata.get('modDate', ''),
            'criptografado': pdf.is_encrypted,
        }
        
        # Obtém dimensões da primeira página
        if pdf.page_count > 0:
            page = pdf[0]
            rect = page.rect
            info['largura_mm'] = round(rect.width * 25.4 / 72, 1)
            info['altura_mm'] = round(rect.height * 25.4 / 72, 1)
        
        pdf.close()
        return info
        
    except Exception as e:
        logger.error(f"Erro ao obter info do PDF: {e}")
        return None


def _formatar_tamanho(bytes_size: int) -> str:
    """Formata tamanho em bytes para formato legível"""
    if bytes_size < 1024:
        return f"{bytes_size} B"
    elif bytes_size < 1024 * 1024:
        return f"{bytes_size / 1024:.1f} KB"
    else:
        return f"{bytes_size / (1024 * 1024):.1f} MB"


def criar_relatorio_pdf(titulo: str, dados: dict, tipo: str = "financeiro") -> Optional[bytes]:
    """Cria um relatório formatado em PDF
    
    Args:
        titulo: Título do relatório
        dados: Dados do relatório
        tipo: Tipo de relatório ('financeiro', 'vencimentos', etc)
    
    Returns:
        bytes do PDF ou None se falhar
    """
    if not HAS_REPORTLAB:
        logger.error("ReportLab não instalado")
        return None
    
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        elementos = []
        
        # Cabeçalho
        titulo_style = ParagraphStyle(
            'Titulo',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=20,
            textColor=HexColor('#1a5276'),
            alignment=1
        )
        elementos.append(Paragraph(titulo, titulo_style))
        
        # Data do relatório
        data_style = ParagraphStyle(
            'Data',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#666666'),
            alignment=1
        )
        data_atual = datetime.now().strftime('%d/%m/%Y às %H:%M')
        elementos.append(Paragraph(f"Gerado em {data_atual}", data_style))
        elementos.append(Spacer(1, 30))
        
        if tipo == "financeiro":
            elementos.extend(_criar_secao_financeiro(dados, styles))
        elif tipo == "vencimentos":
            elementos.extend(_criar_secao_vencimentos(dados, styles))
        else:
            # Genérico - lista os dados
            for chave, valor in dados.items():
                elementos.append(Paragraph(f"<b>{chave}:</b> {valor}", styles['Normal']))
                elementos.append(Spacer(1, 5))
        
        # Rodapé
        elementos.append(Spacer(1, 40))
        rodape_style = ParagraphStyle(
            'Rodape',
            parent=styles['Normal'],
            fontSize=8,
            textColor=HexColor('#999999'),
            alignment=1
        )
        elementos.append(Paragraph("Relatório gerado automaticamente pelo Assistente Ranny", rodape_style))
        
        doc.build(elementos)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"Relatório PDF criado: {len(result)} bytes")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar relatório PDF: {e}")
        return None


def _criar_secao_financeiro(dados: dict, styles) -> list:
    """Cria seção de relatório financeiro"""
    elementos = []
    
    # Resumo
    if 'resumo' in dados:
        subtitulo = ParagraphStyle(
            'Subtitulo',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=HexColor('#2c3e50'),
            spaceAfter=10
        )
        elementos.append(Paragraph("📊 Resumo", subtitulo))
        
        resumo = dados['resumo']
        texto_resumo = f"""
        <b>Total:</b> R$ {resumo.get('total', 0):,.2f}<br/>
        <b>Média diária:</b> R$ {resumo.get('media', 0):,.2f}<br/>
        <b>Período:</b> {resumo.get('periodo', 'N/A')}
        """
        elementos.append(Paragraph(texto_resumo, styles['Normal']))
        elementos.append(Spacer(1, 20))
    
    # Tabela de fechamentos
    if 'fechamentos' in dados and dados['fechamentos']:
        elementos.append(Paragraph("💰 Fechamentos", styles['Heading2']))
        elementos.append(Spacer(1, 10))
        
        # Cabeçalho da tabela
        table_data = [['Data', 'Valor']]
        
        for f in dados['fechamentos'][:30]:  # Limita a 30 registros
            data = f.get('data', '')[:10] if f.get('data') else ''
            valor = f"R$ {f.get('valor', 0):,.2f}"
            table_data.append([data, valor])
        
        table = Table(table_data, colWidths=[8*cm, 6*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#3498db')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), HexColor('#ecf0f1')),
            ('GRID', (0, 0), (-1, -1), 1, HexColor('#bdc3c7')),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#ffffff'), HexColor('#f8f9fa')]),
        ]))
        elementos.append(table)
    
    return elementos


def _criar_secao_vencimentos(dados: dict, styles) -> list:
    """Cria seção de relatório de vencimentos"""
    elementos = []
    
    if 'vencimentos' in dados and dados['vencimentos']:
        # Tabela de vencimentos
        table_data = [['Descrição', 'Vencimento', 'Valor', 'Status']]
        
        for v in dados['vencimentos']:
            desc = v.get('descricao', '')[:30]
            venc = v.get('vencimento', '')[:10] if v.get('vencimento') else ''
            valor = f"R$ {v.get('valor', 0):,.2f}" if v.get('valor') else '-'
            status = '✅ Pago' if v.get('pago') else '⏳ Pendente'
            table_data.append([desc, venc, valor, status])
        
        table = Table(table_data, colWidths=[6*cm, 3*cm, 3*cm, 3*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#e74c3c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, HexColor('#bdc3c7')),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#ffffff'), HexColor('#fdf2f2')]),
        ]))
        elementos.append(table)
    
    return elementos


# ============ FUNÇÕES PARA DOCX (Word) ============

def criar_docx_texto(texto: str, titulo: str = None) -> Optional[bytes]:
    """Cria um documento Word (DOCX) a partir de texto
    
    Args:
        texto: Conteúdo do documento
        titulo: Título opcional do documento
    
    Returns:
        bytes do DOCX ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document()
        
        # Adiciona título se fornecido
        if titulo:
            heading = doc.add_heading(titulo, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Espaço após título
        
        # Processa o texto - divide em parágrafos
        paragrafos = texto.split('\n')
        for p in paragrafos:
            p = p.strip()
            if p:
                # Verifica se é um item de lista numerada (1. ou 1) ou com símbolos (• ou -)
                # CORREÇÃO #1: Adiciona suporte a listas numeradas
                if re.match(r'^\d+[.\)]\s', p):
                    # Lista numerada (1. ou 1) )
                    item_texto = re.sub(r'^\d+[.\)]\s', '', p).strip()
                    doc.add_paragraph(item_texto, style='List Number')
                elif p.startswith('•') or p.startswith('-'):
                    # Lista com símbolos (• ou -)
                    item_texto = p.lstrip('•-').strip()
                    doc.add_paragraph(item_texto, style='List Bullet')
                else:
                    doc.add_paragraph(p)
            else:
                doc.add_paragraph()  # Linha em branco
        
        # Adiciona rodapé com data
        doc.add_paragraph()
        rodape = doc.add_paragraph()
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rodape.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')} • Assistente Ranny")
        run.font.size = Pt(8)
        run.font.italic = True
        
        # Salva em bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"DOCX criado com sucesso: {len(result)} bytes")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar DOCX: {e}")
        return None


def criar_docx_tabela(dados: list, cabecalho: list = None, titulo: str = None) -> Optional[bytes]:
    """Cria um documento Word com uma tabela
    
    Args:
        dados: Lista de listas com os dados da tabela
        cabecalho: Lista com os títulos das colunas (opcional)
        titulo: Título do documento (opcional)
    
    Returns:
        bytes do DOCX ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document()
        
        # Adiciona título se fornecido
        if titulo:
            heading = doc.add_heading(titulo, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Determina número de colunas
        if cabecalho:
            num_cols = len(cabecalho)
        elif dados:
            num_cols = len(dados[0]) if isinstance(dados[0], (list, tuple)) else 1
        else:
            logger.error("Dados vazios para criar tabela")
            return None
        
        # Cria tabela
        num_rows = len(dados) + (1 if cabecalho else 0)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Adiciona cabeçalho
        row_idx = 0
        if cabecalho:
            for col_idx, header in enumerate(cabecalho):
                cell = table.rows[0].cells[col_idx]
                cell.text = str(header)
                # Negrito no cabeçalho
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            row_idx = 1
        
        # Adiciona dados
        for data_row in dados:
            if isinstance(data_row, (list, tuple)):
                for col_idx, value in enumerate(data_row):
                    if col_idx < num_cols:
                        table.rows[row_idx].cells[col_idx].text = str(value)
            else:
                table.rows[row_idx].cells[0].text = str(data_row)
            row_idx += 1
        
        # Rodapé
        doc.add_paragraph()
        rodape = doc.add_paragraph()
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rodape.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')} • Assistente Ranny")
        run.font.size = Pt(8)
        run.font.italic = True
        
        # Salva em bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"DOCX com tabela criado: {len(result)} bytes")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar DOCX com tabela: {e}")
        return None


# ============ FUNÇÕES PARA XLSX (Excel) ============

def criar_xlsx_texto(texto: str, titulo: str = None) -> Optional[bytes]:
    """Cria uma planilha Excel (XLSX) a partir de texto
    
    O texto é dividido em linhas e cada linha vira uma célula na coluna A.
    Se houver separadores (vírgula, ponto-e-vírgula, tab), divide em colunas.
    
    Args:
        texto: Conteúdo da planilha
        titulo: Título opcional (vai na primeira linha)
    
    Returns:
        bytes do XLSX ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Dados"
        
        # Estilos
        header_font = Font(bold=True, size=14, color="FFFFFF")
        header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center")
        
        row_num = 1
        
        # Adiciona título se fornecido
        if titulo:
            ws.cell(row=row_num, column=1, value=titulo)
            ws.cell(row=row_num, column=1).font = header_font
            ws.cell(row=row_num, column=1).fill = header_fill
            ws.cell(row=row_num, column=1).alignment = header_align
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=5)
            row_num += 2  # Pula uma linha após o título
        
        # Processa o texto
        linhas = texto.split('\n')
        for linha in linhas:
            linha = linha.strip()
            if linha:
                # Remove bullet points
                if linha.startswith('•') or linha.startswith('-'):
                    linha = linha.lstrip('•-').strip()
                
                # Tenta detectar separador
                if '\t' in linha:
                    valores = linha.split('\t')
                elif ';' in linha:
                    valores = linha.split(';')
                elif ',' in linha and linha.count(',') >= 2:
                    valores = linha.split(',')
                else:
                    valores = [linha]
                
                for col_num, valor in enumerate(valores, 1):
                    ws.cell(row=row_num, column=col_num, value=valor.strip())
                
                row_num += 1
        
        # Ajusta largura das colunas
        for col in range(1, ws.max_column + 1):
            max_length = 0
            column_letter = get_column_letter(col)
            for row in range(1, ws.max_row + 1):
                cell = ws.cell(row=row, column=col)
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
        
        # Salva em bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"XLSX criado com sucesso: {len(result)} bytes")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar XLSX: {e}")
        return None


def criar_xlsx_tabela(dados: list, cabecalho: list = None, titulo: str = None) -> Optional[bytes]:
    """Cria uma planilha Excel com dados tabulares
    
    Args:
        dados: Lista de listas com os dados
        cabecalho: Lista com os títulos das colunas (opcional)
        titulo: Título da planilha (opcional)
    
    Returns:
        bytes do XLSX ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = titulo[:31] if titulo else "Dados"  # Excel limita a 31 caracteres
        
        # Estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center")
        
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        row_num = 1
        
        # Adiciona cabeçalho
        if cabecalho:
            for col_num, header in enumerate(cabecalho, 1):
                cell = ws.cell(row=row_num, column=col_num, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                cell.border = thin_border
            row_num += 1
        
        # Adiciona dados
        alt_fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
        
        for idx, data_row in enumerate(dados):
            if isinstance(data_row, (list, tuple)):
                for col_num, value in enumerate(data_row, 1):
                    cell = ws.cell(row=row_num, column=col_num, value=value)
                    cell.border = thin_border
                    if idx % 2 == 1:  # Linhas alternadas
                        cell.fill = alt_fill
            else:
                cell = ws.cell(row=row_num, column=1, value=data_row)
                cell.border = thin_border
                if idx % 2 == 1:
                    cell.fill = alt_fill
            row_num += 1
        
        # Ajusta largura das colunas
        for col in range(1, ws.max_column + 1):
            max_length = 0
            column_letter = get_column_letter(col)
            for row in range(1, ws.max_row + 1):
                cell = ws.cell(row=row, column=col)
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
        
        # Salva em bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"XLSX com tabela criado: {len(result)} bytes")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar XLSX com tabela: {e}")
        return None


def criar_xlsx_lista(itens: list, titulo: str = None, cabecalho: str = "Item") -> Optional[bytes]:
    """Cria uma planilha Excel simples com uma lista de itens
    
    Args:
        itens: Lista de itens
        titulo: Título da planilha (opcional)
        cabecalho: Título da coluna (padrão: "Item")
    
    Returns:
        bytes do XLSX ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        # Converte para formato de tabela
        dados = [[item] for item in itens]
        return criar_xlsx_tabela(dados, cabecalho=[cabecalho], titulo=titulo)
        
    except Exception as e:
        logger.error(f"Erro ao criar XLSX de lista: {e}")
        return None


# ============ FUNÇÕES DE LEITURA ============

def ler_docx(docx_bytes: bytes) -> Optional[dict]:
    """Lê um documento Word e retorna seu conteúdo estruturado

    Args:
        docx_bytes: bytes do arquivo DOCX

    Returns:
        dict com 'texto', 'paragrafos', 'tabelas', 'headers', 'footers' ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None

    try:
        doc = Document(io.BytesIO(docx_bytes))

        # Extrai parágrafos
        paragrafos = []
        texto_completo = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragrafos.append({
                    'texto': para.text,
                    'estilo': para.style.name if para.style else 'Normal'
                })
                texto_completo.append(para.text)

        # Extrai tabelas
        tabelas = []
        for table in doc.tables:
            dados_tabela = []
            for row in table.rows:
                linha = [cell.text for cell in row.cells]
                dados_tabela.append(linha)
            tabelas.append(dados_tabela)

        # CORREÇÃO #4: Extrai headers e footers
        headers = []
        footers = []

        # Itera sobre as seções do documento
        for section in doc.sections:
            # Headers
            header = section.header
            if header.is_linked_to_previous == False:
                # Header próprio da seção
                for para in header.paragraphs:
                    if para.text.strip():
                        headers.append(para.text)

            # Footers
            footer = section.footer
            if footer.is_linked_to_previous == False:
                # Footer próprio da seção
                for para in footer.paragraphs:
                    if para.text.strip():
                        footers.append(para.text)

        return {
            'texto': '\n'.join(texto_completo),
            'paragrafos': paragrafos,
            'tabelas': tabelas,
            'num_paragrafos': len(paragrafos),
            'num_tabelas': len(tabelas),
            'headers': headers,  # CORREÇÃO #4
            'footers': footers,  # CORREÇÃO #4
            'num_headers': len(headers),
            'num_footers': len(footers)
        }

    except Exception as e:
        logger.error(f"Erro ao ler DOCX: {e}")
        return None


def ler_xlsx(xlsx_bytes: bytes) -> Optional[dict]:
    """Lê uma planilha Excel e retorna seu conteúdo estruturado
    
    Args:
        xlsx_bytes: bytes do arquivo XLSX
    
    Returns:
        dict com 'planilhas', 'texto' ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = load_workbook(io.BytesIO(xlsx_bytes), read_only=True)
        
        planilhas = {}
        texto_completo = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            dados = []
            for row in sheet.iter_rows(values_only=True):
                linha = [str(cell) if cell is not None else '' for cell in row]
                if any(linha):  # Ignora linhas vazias
                    dados.append(linha)
                    texto_completo.append(' | '.join(linha))
            planilhas[sheet_name] = dados
        
        wb.close()
        
        return {
            'planilhas': planilhas,
            'texto': '\n'.join(texto_completo),
            'nomes_planilhas': list(planilhas.keys()),
            'num_planilhas': len(planilhas)
        }
        
    except Exception as e:
        logger.error(f"Erro ao ler XLSX: {e}")
        return None


# ============ FUNÇÕES DE EDIÇÃO DOCX ============

def editar_docx_adicionar_texto(docx_bytes: bytes, texto: str, posicao: str = 'fim') -> Optional[bytes]:
    """Adiciona texto a um documento Word existente
    
    Args:
        docx_bytes: bytes do arquivo DOCX original
        texto: texto a adicionar
        posicao: 'inicio' ou 'fim' (padrão: 'fim')
    
    Returns:
        bytes do DOCX modificado ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        # Processa o texto (converte listas com " - " em parágrafos)
        linhas = texto.replace(' - ', '\n• ').split('\n')
        
        if posicao == 'inicio':
            # Adiciona no início (antes do primeiro parágrafo)
            for i, linha in enumerate(reversed(linhas)):
                if linha.strip():
                    para = doc.paragraphs[0].insert_paragraph_before(linha.strip())
        else:
            # Adiciona no fim
            for linha in linhas:
                if linha.strip():
                    doc.add_paragraph(linha.strip())
        
        # Salva em bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao editar DOCX (adicionar texto): {e}")
        return None


def editar_docx_substituir(docx_bytes: bytes, texto_antigo: str, texto_novo: str) -> Optional[Tuple[bytes, int]]:
    """Substitui texto em um documento Word (case-insensitive)

    Args:
        docx_bytes: bytes do arquivo DOCX original
        texto_antigo: texto a ser substituído
        texto_novo: novo texto

    Returns:
        Tuple (bytes do DOCX modificado, número de substituições) ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None

    try:
        doc = Document(io.BytesIO(docx_bytes))
        substituicoes = 0

        # CORREÇÃO #2: Usa padrão case-insensitive
        # CORREÇÃO #3: Lida melhor com runs divididos
        texto_antigo_lower = texto_antigo.lower()

        # Substitui em parágrafos
        for para in doc.paragraphs:
            texto_para = para.text
            if texto_antigo_lower in texto_para.lower():
                # Conta quantas substituições serão feitas
                count = texto_para.lower().count(texto_antigo_lower)
                substituicoes += count

                # Se há múltiplos runs e texto está dividido, reconstruí o parágrafo
                if len(para.runs) > 1:
                    # Verifica se o texto está todo junto ou dividido
                    texto_total = texto_para.lower()
                    if texto_antigo_lower in texto_total:
                        # Texto está dividido em runs - usa approach de reconstruir
                        # Primeiro, tenta substituição simples em cada run
                        novo_texto = re.sub(re.escape(texto_antigo), texto_novo, texto_para, flags=re.IGNORECASE)
                        # Limpa todos os runs e recria com o texto novo
                        for run in para.runs:
                            run.text = ''
                        if para.runs:
                            para.runs[0].text = novo_texto
                else:
                    # Runs únicos - substituição direta
                    for run in para.runs:
                        if texto_antigo_lower in run.text.lower():
                            # Preserva a formatação do primeiro character do run
                            run.text = re.sub(re.escape(texto_antigo), texto_novo, run.text, flags=re.IGNORECASE)

        # Substitui em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_cell = cell.text
                    if texto_antigo_lower in texto_cell.lower():
                        count = texto_cell.lower().count(texto_antigo_lower)
                        substituicoes += count

                        for para in cell.paragraphs:
                            if texto_antigo_lower in para.text.lower():
                                if len(para.runs) > 1:
                                    novo_texto = re.sub(re.escape(texto_antigo), texto_novo, para.text, flags=re.IGNORECASE)
                                    for run in para.runs:
                                        run.text = ''
                                    if para.runs:
                                        para.runs[0].text = novo_texto
                                else:
                                    for run in para.runs:
                                        if texto_antigo_lower in run.text.lower():
                                            run.text = re.sub(re.escape(texto_antigo), texto_novo, run.text, flags=re.IGNORECASE)

        # Salva em bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue(), substituicoes

    except Exception as e:
        logger.error(f"Erro ao editar DOCX (substituir): {e}")
        return None


def editar_docx_remover_paragrafo(docx_bytes: bytes, texto_busca: str) -> Optional[Tuple[bytes, int]]:
    """Remove parágrafos que contêm determinado texto
    
    Args:
        docx_bytes: bytes do arquivo DOCX original
        texto_busca: texto a buscar (parágrafos contendo serão removidos)
    
    Returns:
        Tuple (bytes do DOCX modificado, número de remoções) ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        removidos = 0
        
        # Identifica parágrafos a remover
        for para in doc.paragraphs:
            if texto_busca.lower() in para.text.lower():
                # Remove o parágrafo
                p = para._element
                p.getparent().remove(p)
                removidos += 1
        
        # Salva em bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue(), removidos
        
    except Exception as e:
        logger.error(f"Erro ao editar DOCX (remover): {e}")
        return None


def duplicar_docx(docx_bytes: bytes, novo_nome: str = None) -> Optional[bytes]:
    """Duplica um documento DOCX com novo nome

    Args:
        docx_bytes: bytes do arquivo DOCX original
        novo_nome: novo nome para o documento (opcional, apenas para logging)

    Returns:
        bytes do DOCX duplicado ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None

    try:
        # Apenas abre e salva novamente (docx é imutável como zip)
        doc = Document(io.BytesIO(docx_bytes))

        # Salva em bytes
        output = io.BytesIO()
        doc.save(output)

        logger.info(f"DOCX duplicado: {novo_nome or 'sem nome'}")
        return output.getvalue()

    except Exception as e:
        logger.error(f"Erro ao duplicar DOCX: {e}")
        return None


# ============ FUNÇÕES DE EDIÇÃO XLSX ============

def editar_xlsx_adicionar_linha(xlsx_bytes: bytes, dados: list, planilha: str = None) -> Optional[bytes]:
    """Adiciona uma linha a uma planilha Excel
    
    Args:
        xlsx_bytes: bytes do arquivo XLSX original
        dados: lista com os valores da nova linha
        planilha: nome da planilha (usa a primeira se não especificado)
    
    Returns:
        bytes do XLSX modificado ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = load_workbook(io.BytesIO(xlsx_bytes))
        
        # Seleciona a planilha
        if planilha and planilha in wb.sheetnames:
            ws = wb[planilha]
        else:
            ws = wb.active
        
        # Adiciona a linha no final com estilo
        next_row = ws.max_row + 1
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        data_align = Alignment(horizontal="left", vertical="center")
        
        for col_num, valor in enumerate(dados, 1):
            cell = ws.cell(row=next_row, column=col_num, value=valor)
            cell.border = thin_border
            cell.alignment = data_align
        
        # Salva em bytes
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao editar XLSX (adicionar linha): {e}")
        return None


def editar_xlsx_adicionar_linhas(xlsx_bytes: bytes, linhas: list, planilha: str = None) -> Optional[bytes]:
    """Adiciona múltiplas linhas a uma planilha Excel
    
    Args:
        xlsx_bytes: bytes do arquivo XLSX original
        linhas: lista de listas com os valores das novas linhas
        planilha: nome da planilha (usa a primeira se não especificado)
    
    Returns:
        bytes do XLSX modificado ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = load_workbook(io.BytesIO(xlsx_bytes))
        
        # Seleciona a planilha
        if planilha and planilha in wb.sheetnames:
            ws = wb[planilha]
        else:
            ws = wb.active
        
        # Adiciona as linhas com estilo
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        data_align = Alignment(horizontal="left", vertical="center")
        
        for linha in linhas:
            next_row = ws.max_row + 1
            if isinstance(linha, (list, tuple)):
                for col_num, valor in enumerate(linha, 1):
                    cell = ws.cell(row=next_row, column=col_num, value=valor)
                    cell.border = thin_border
                    cell.alignment = data_align
            else:
                cell = ws.cell(row=next_row, column=1, value=linha)
                cell.border = thin_border
                cell.alignment = data_align
        
        # Salva em bytes
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao editar XLSX (adicionar linhas): {e}")
        return None


def editar_xlsx_substituir(xlsx_bytes: bytes, texto_antigo: str, texto_novo: str, planilha: str = None) -> Optional[Tuple[bytes, int]]:
    """Substitui texto em uma planilha Excel
    
    Args:
        xlsx_bytes: bytes do arquivo XLSX original
        texto_antigo: texto a ser substituído
        texto_novo: novo texto
        planilha: nome da planilha (todas se não especificado)
    
    Returns:
        Tuple (bytes do XLSX modificado, número de substituições) ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = load_workbook(io.BytesIO(xlsx_bytes))
        substituicoes = 0
        
        # Define quais planilhas processar
        if planilha and planilha in wb.sheetnames:
            sheets = [wb[planilha]]
        else:
            sheets = wb.worksheets
        
        # Processa cada planilha
        for ws in sheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str) and texto_antigo in cell.value:
                        substituicoes += cell.value.count(texto_antigo)
                        cell.value = cell.value.replace(texto_antigo, texto_novo)
        
        # Salva em bytes
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue(), substituicoes
        
    except Exception as e:
        logger.error(f"Erro ao editar XLSX (substituir): {e}")
        return None


def editar_xlsx_remover_linha(xlsx_bytes: bytes, numero_linha: int, planilha: str = None) -> Optional[bytes]:
    """Remove uma linha específica de uma planilha Excel
    
    Args:
        xlsx_bytes: bytes do arquivo XLSX original
        numero_linha: número da linha a remover (1-indexed)
        planilha: nome da planilha (usa a primeira se não especificado)
    
    Returns:
        bytes do XLSX modificado ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = load_workbook(io.BytesIO(xlsx_bytes))
        
        # Seleciona a planilha
        if planilha and planilha in wb.sheetnames:
            ws = wb[planilha]
        else:
            ws = wb.active
        
        # Validação de limites
        if numero_linha < 1:
            logger.error("Número de linha inválido (< 1)")
            return None
        
        if numero_linha > ws.max_row:
            logger.error(f"Linha {numero_linha} não existe (máx: {ws.max_row})")
            return None
        
        # Remove a linha
        ws.delete_rows(numero_linha)
        
        # Salva em bytes
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao editar XLSX (remover linha): {e}")
        return None


def editar_xlsx_atualizar_celula(xlsx_bytes: bytes, linha: int, coluna: int, valor, planilha: str = None) -> Optional[bytes]:
    """Atualiza o valor de uma célula específica
    
    Args:
        xlsx_bytes: bytes do arquivo XLSX original
        linha: número da linha (1-indexed)
        coluna: número da coluna (1-indexed) ou letra (A, B, C...)
        valor: novo valor da célula
        planilha: nome da planilha (usa a primeira se não especificado)
    
    Returns:
        bytes do XLSX modificado ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = load_workbook(io.BytesIO(xlsx_bytes))
        
        # Seleciona a planilha
        if planilha and planilha in wb.sheetnames:
            ws = wb[planilha]
        else:
            ws = wb.active
        
        # Converte coluna letra para número se necessário
        if isinstance(coluna, str):
            from openpyxl.utils import column_index_from_string
            coluna = column_index_from_string(coluna)
        
        # Atualiza a célula
        ws.cell(row=linha, column=coluna, value=valor)
        
        # Salva em bytes
        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao editar XLSX (atualizar célula): {e}")
        return None


def criar_xlsx_entregadores(dados: dict, custo_semana: float = 1.0, custo_fds: float = 10.0, 
                           bonus_horario: float = 10.0, custo_entrega: float = 12.0) -> Optional[bytes]:
    """Cria planilha Excel formatada para controle de entregadores
    
    Args:
        dados: Dicionário com estrutura:
            {
                "periodo": "Semana 10/02 a 16/02",
                "dias": [
                    {"dia": "segunda", "entregadores": 3, "chegaram_horario": 0, "entregas": 20},
                    ...
                ]
            }
        custo_semana: Custo por entregador seg-qui (padrão: 1.0)
        custo_fds: Custo por entregador sex-dom (padrão: 10.0)
        bonus_horario: Bônus por chegar até 18:10 no FDS (padrão: 10.0)
        custo_entrega: Custo por entrega (padrão: 12.0)
    
    Returns:
        bytes do XLSX ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Entregadores"
        
        # Estilos
        title_font = Font(bold=True, size=14, color="FFFFFF")
        title_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        title_align = Alignment(horizontal="center", vertical="center")
        
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        
        data_align = Alignment(horizontal="center", vertical="center")
        currency_align = Alignment(horizontal="right", vertical="center")
        
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        thick_border = Border(
            left=Side(style='medium'),
            right=Side(style='medium'),
            top=Side(style='medium'),
            bottom=Side(style='medium')
        )
        
        # Linha 1: Título
        ws.merge_cells('A1:H1')
        title_cell = ws['A1']
        title_cell.value = f"📊 CONTROLE DE ENTREGADORES - {dados.get('periodo', 'Semana')}"
        title_cell.font = title_font
        title_cell.fill = title_fill
        title_cell.alignment = title_align
        title_cell.border = thick_border
        ws.row_dimensions[1].height = 30
        
        # Linha 2: Vazia (espaçamento)
        ws.row_dimensions[2].height = 5
        
        # Linha 3: Cabeçalhos
        headers = [
            "Dia",
            "Entregadores",
            "Chegaram\n18:10",
            "Entregas",
            "Custo\nEntregadores",
            "Bônus\nHorário",
            "Custo\nEntregas",
            "TOTAL"
        ]
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col_num, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
        
        ws.row_dimensions[3].height = 35
        
        # Dados dos dias
        dias_data = dados.get('dias', [])
        row_num = 4
        
        # Detecta formato: se algum dia contém "/", é formato de data (mensal)
        formato_data = any('/' in str(dia_info.get('dia', '')) for dia_info in dias_data)
        
        # Mapeamento de dias para identificar fim de semana
        dias_fds = {'sexta', 'sabado', 'sábado', 'domingo'}
        
        # Função auxiliar para determinar se é FDS
        def is_fim_de_semana(dia_str):
            if formato_data:
                # Formato de data: parseia e verifica weekday
                try:
                    from datetime import datetime
                    # Assume formato DD/MM
                    dia_num, mes_num = dia_str.split('/')
                    ano_atual = datetime.now().year
                    data = datetime(ano_atual, int(mes_num), int(dia_num))
                    # weekday: 0=segunda, 4=sexta, 5=sábado, 6=domingo
                    return data.weekday() >= 4
                except:
                    return False
            else:
                # Formato de nome de dia
                return any(d in dia_str.lower() for d in dias_fds)
        
        alt_fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
        fds_fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
        
        for idx, dia_info in enumerate(dias_data):
            dia = dia_info.get('dia', '').lower()
            entregadores_raw = dia_info.get('entregadores', 0)
            chegaram_horario_raw = dia_info.get('chegaram_horario', 0)
            entregas_raw = dia_info.get('entregas', 0)

            # CORREÇÃO #1: Normaliza entregadores para int (pode vir como lista ou string)
            if isinstance(entregadores_raw, list):
                entregadores = len(entregadores_raw)
            elif isinstance(entregadores_raw, str):
                try:
                    entregadores = int(entregadores_raw)
                except ValueError:
                    entregadores = 0
            else:
                entregadores = int(entregadores_raw or 0)

            # CORREÇÃO #1: Normaliza chegaram_horario para int
            if isinstance(chegaram_horario_raw, str):
                try:
                    chegaram_horario = int(chegaram_horario_raw)
                except ValueError:
                    chegaram_horario = 0
            else:
                chegaram_horario = int(chegaram_horario_raw or 0)

            # CORREÇÃO #1: Normaliza entregas para int
            if isinstance(entregas_raw, str):
                try:
                    entregas = int(entregas_raw)
                except ValueError:
                    entregas = 0
            else:
                entregas = int(entregas_raw or 0)

            # Determina se é fim de semana
            is_fds = is_fim_de_semana(dia)
            
            # Formata dia para exibição
            dia_display = dia.capitalize() if not formato_data else dia
            
            # Coluna A: Dia
            cell = ws.cell(row=row_num, column=1, value=dia_display)
            cell.alignment = data_align
            cell.border = thin_border
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            # Coluna B: Entregadores
            cell = ws.cell(row=row_num, column=2, value=entregadores)
            cell.alignment = data_align
            cell.border = thin_border
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            # Coluna C: Chegaram 18:10
            # CORREÇÃO #3: Mostra valor mesmo em dias úteis (documenta para referência)
            # O valor entra no cálculo apenas em FDS
            cell = ws.cell(row=row_num, column=3, value=chegaram_horario if is_fds else ('-' if chegaram_horario == 0 else chegaram_horario))
            cell.alignment = data_align
            cell.border = thin_border
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            # Coluna D: Entregas
            cell = ws.cell(row=row_num, column=4, value=entregas)
            cell.alignment = data_align
            cell.border = thin_border
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            # Coluna E: Custo Entregadores (fórmula)
            if is_fds:
                formula = f"=B{row_num}*{custo_fds}"
            else:
                formula = f"=B{row_num}*{custo_semana}"
            
            cell = ws.cell(row=row_num, column=5, value=formula)
            cell.number_format = 'R$ #,##0.00'
            cell.alignment = currency_align
            cell.border = thin_border
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            # Coluna F: Bônus Horário (fórmula)
            if is_fds:
                formula = f"=C{row_num}*{bonus_horario}"
            else:
                formula = "=0"
            
            cell = ws.cell(row=row_num, column=6, value=formula)
            cell.number_format = 'R$ #,##0.00'
            cell.alignment = currency_align
            cell.border = thin_border
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            # Coluna G: Custo Entregas (fórmula)
            formula = f"=D{row_num}*{custo_entrega}"
            cell = ws.cell(row=row_num, column=7, value=formula)
            cell.number_format = 'R$ #,##0.00'
            cell.alignment = currency_align
            cell.border = thin_border
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            # Coluna H: TOTAL (fórmula)
            formula = f"=E{row_num}+F{row_num}+G{row_num}"
            cell = ws.cell(row=row_num, column=8, value=formula)
            cell.number_format = 'R$ #,##0.00'
            cell.alignment = currency_align
            cell.border = thin_border
            cell.font = Font(bold=True)
            if is_fds:
                cell.fill = fds_fill
            elif idx % 2 == 1:
                cell.fill = alt_fill
            
            row_num += 1
        
        # Linha de TOTAL
        total_row = row_num
        total_fill = PatternFill(start_color="27AE60", end_color="27AE60", fill_type="solid")
        total_font = Font(bold=True, color="FFFFFF", size=12)
        
        # Coluna A: "TOTAL"
        cell = ws.cell(row=total_row, column=1, value="TOTAL")
        cell.font = total_font
        cell.fill = total_fill
        cell.alignment = data_align
        cell.border = thick_border
        
        # Colunas B, C, D: Somas
        for col in [2, 3, 4]:
            formula = f"=SUM({get_column_letter(col)}4:{get_column_letter(col)}{total_row-1})"
            cell = ws.cell(row=total_row, column=col, value=formula)
            cell.font = total_font
            cell.fill = total_fill
            cell.alignment = data_align
            cell.border = thick_border
        
        # Colunas E, F, G, H: Somas com formato moeda
        for col in [5, 6, 7, 8]:
            formula = f"=SUM({get_column_letter(col)}4:{get_column_letter(col)}{total_row-1})"
            cell = ws.cell(row=total_row, column=col, value=formula)
            cell.number_format = 'R$ #,##0.00'
            cell.font = total_font
            cell.fill = total_fill
            cell.alignment = currency_align
            cell.border = thick_border
        
        ws.row_dimensions[total_row].height = 25
        
        # Ajusta largura das colunas
        column_widths = {
            'A': 12,  # Dia
            'B': 12,  # Entregadores
            'C': 12,  # Chegaram 18:10
            'D': 10,  # Entregas
            'E': 15,  # Custo Entregadores
            'F': 15,  # Bônus Horário
            'G': 15,  # Custo Entregas
            'H': 18   # TOTAL
        }
        
        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width
        
        # ===== ADICIONA GRÁFICOS =====
        try:
            from openpyxl.chart import BarChart, PieChart, Reference
            from openpyxl.chart.label import DataLabelList
            
            # GRÁFICO 1: Entregas por Dia (Colunas)
            chart1 = BarChart()
            chart1.type = "col"
            chart1.style = 10
            chart1.title = "📊 Entregas por Dia"
            chart1.y_axis.title = 'Quantidade'
            chart1.x_axis.title = 'Dia'
            
            # Dados: Dias (A4:A{total_row-1}) e Entregas (D4:D{total_row-1})
            data1 = Reference(ws, min_col=4, min_row=3, max_row=total_row-1)
            cats1 = Reference(ws, min_col=1, min_row=4, max_row=total_row-1)
            chart1.add_data(data1, titles_from_data=True)
            chart1.set_categories(cats1)
            
            # Estilo
            chart1.height = 10
            chart1.width = 20
            
            # Posiciona o gráfico
            ws.add_chart(chart1, f"J3")
            
            # GRÁFICO 2: Distribuição de Custos (Pizza)
            chart2 = PieChart()
            chart2.title = "💰 Distribuição de Custos"
            
            # Dados: Totais de cada tipo de custo (E{total_row}, F{total_row}, G{total_row})
            # Cria referências para os totais
            data2 = Reference(ws, min_col=5, min_row=total_row, max_col=7, max_row=total_row)
            labels2 = Reference(ws, min_col=5, min_row=3, max_col=7, max_row=3)
            
            chart2.add_data(data2, titles_from_data=True)
            chart2.set_categories(labels2)
            
            # Mostra valores e porcentagens
            chart2.dataLabels = DataLabelList()
            chart2.dataLabels.showPercent = True
            chart2.dataLabels.showVal = True
            
            # Estilo
            chart2.height = 10
            chart2.width = 15
            
            # Posiciona o gráfico
            ws.add_chart(chart2, f"J19")
            
            # GRÁFICO 3: Custo Total por Dia (Barras Horizontais)
            chart3 = BarChart()
            chart3.type = "bar"
            chart3.style = 11
            chart3.title = "💵 Custo Total por Dia"
            chart3.y_axis.title = 'Dia'
            chart3.x_axis.title = 'Valor (R$)'
            
            # Dados: Dias e Total (H4:H{total_row-1})
            data3 = Reference(ws, min_col=8, min_row=3, max_row=total_row-1)
            cats3 = Reference(ws, min_col=1, min_row=4, max_row=total_row-1)
            chart3.add_data(data3, titles_from_data=True)
            chart3.set_categories(cats3)
            
            # Estilo
            chart3.height = 10
            chart3.width = 20
            
            # Posiciona o gráfico
            ws.add_chart(chart3, f"W3")
            
            logger.info("✅ Gráficos adicionados à planilha")
            
        except ImportError:
            logger.warning("⚠️ openpyxl.chart não disponível - gráficos não adicionados")
        except Exception as e:
            logger.warning(f"⚠️ Erro ao adicionar gráficos: {e}")
        
        # Salva em bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"XLSX de entregadores criado: {len(result)} bytes")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar XLSX de entregadores: {e}")
        return None


def criar_xlsx_entregadores_com_nomes(dados: dict, entregadores_fixos: list = None, 
                                      custo_entrega: float = 12.0) -> Optional[bytes]:
    """Cria planilha Excel com nomes dos entregadores (Versão 2 - para Ranny)
    
    Estrutura: Linhas = Entregadores, Colunas = Dias da SEMANA informada
    
    Args:
        dados: Dicionário com estrutura:
            {
                "periodo": "Semana 10/02 a 16/02",
                "dias": [
                    {"dia": "segunda", "entregadores": ["João", "Pedro"], "entregas": 20},
                    ...
                ]
            }
        entregadores_fixos: Lista de nomes dos entregadores fixos (opcional)
        custo_entrega: Custo por entrega (padrão: 12.0)
    
    Returns:
        bytes do XLSX ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        from datetime import datetime
        
        # Lista padrão de entregadores fixos
        # CORREÇÃO #5: Agora usa config.ENTREGADORES_FIXOS se não especificado
        if entregadores_fixos is None:
            try:
                from config import ENTREGADORES_FIXOS
                entregadores_fixos = ENTREGADORES_FIXOS
            except ImportError:
                logger.warning("config.ENTREGADORES_FIXOS não encontrado, usando padrão")
                entregadores_fixos = [
                    "Maycon", "Gustavo Campos", "Gustavo Henrique", "Leonardo",
                    "Sidnei", "Maurício", "Iago", "João Pedro", "José", "Davi",
                    "Ryan", "Kaique", "Brayan"
                ]

        # Extrai todos os nomes mencionados nos dias
        todos_nomes = set()
        entregas_por_dia_pessoa = {}  # {dia: {nome: num_entregas}}
        
        dias_data = dados.get('dias', [])
        
        # Detecta formato: se algum dia contém "/", é formato de data (mensal)
        formato_data = any('/' in str(dia_info.get('dia', '')) for dia_info in dias_data)
        
        # Mapeamento de dias da semana (para formato semanal)
        dias_semana_map = {
            'segunda': 'Segunda',
            'terca': 'Terça',
            'terça': 'Terça',
            'quarta': 'Quarta',
            'quinta': 'Quinta',
            'sexta': 'Sexta',
            'sabado': 'Sábado',
            'sábado': 'Sábado',
            'domingo': 'Domingo'
        }
        
        # Lista ordenada de dias da semana (para formato semanal)
        ordem_dias = ['segunda', 'terca', 'terça', 'quarta', 'quinta', 'sexta', 'sabado', 'sábado', 'domingo']
        dias_presentes = []
        
        for dia_info in dias_data:
            dia = dia_info.get('dia', '').lower()
            nomes = dia_info.get('entregadores', [])
            total_entregas = dia_info.get('entregas', 0)
            
            # Se entregadores é número, converte para lista genérica
            if isinstance(nomes, int):
                nomes = [f"Entregador {i+1}" for i in range(nomes)]
            
            todos_nomes.update(nomes)
            
            # Adiciona dia à lista de dias presentes
            if dia not in dias_presentes:
                dias_presentes.append(dia)
            
            # Distribui entregas igualmente entre entregadores do dia
            if nomes and total_entregas > 0:
                entregas_por_pessoa = total_entregas // len(nomes)
                resto = total_entregas % len(nomes)
                
                entregas_por_dia_pessoa[dia] = {}
                for idx, nome in enumerate(nomes):
                    # Primeiros recebem +1 se houver resto
                    entregas = entregas_por_pessoa + (1 if idx < resto else 0)
                    entregas_por_dia_pessoa[dia][nome] = entregas
        
        # Ordena dias presentes
        if formato_data:
            # Para datas, ordena cronologicamente
            try:
                dias_ordenados = sorted(dias_presentes, key=lambda d: (int(d.split('/')[1]), int(d.split('/')[0])))
            except:
                dias_ordenados = dias_presentes
        else:
            # Para nomes de dias, ordena pela ordem da semana
            dias_ordenados = []
            for dia_ordem in ordem_dias:
                if dia_ordem in dias_presentes:
                    dias_ordenados.append(dia_ordem)
        
        if not dias_ordenados:
            logger.error("Nenhum dia encontrado nos dados")
            return None
        
        # Separa fixos e freelancers
        fixos_normalizados = {nome.lower().strip() for nome in entregadores_fixos}
        
        nomes_fixos = []
        nomes_freelancers = []
        
        for nome in todos_nomes:
            nome_normalizado = nome.lower().strip()
            # Verifica se é fixo (match parcial para variações de nome)
            is_fixo = any(fixo in nome_normalizado or nome_normalizado in fixo 
                         for fixo in fixos_normalizados)
            
            if is_fixo:
                nomes_fixos.append(nome)
            else:
                nomes_freelancers.append(nome)
        
        # Ordena alfabeticamente
        nomes_fixos.sort()
        nomes_freelancers.sort()
        
        # Lista final: fixos + freelancers
        lista_entregadores = nomes_fixos + nomes_freelancers
        
        if not lista_entregadores:
            logger.error("Nenhum entregador encontrado nos dados")
            return None
        
        # Cria workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Entregadores"
        
        # Estilos
        title_font = Font(bold=True, size=14, color="000000")
        title_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        title_align = Alignment(horizontal="center", vertical="center")
        
        header_font = Font(bold=True, size=10, color="000000")
        header_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        
        data_align = Alignment(horizontal="center", vertical="center")
        
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Extrai período
        periodo = dados.get('periodo', 'Semana')
        
        # Linha 1: Título
        num_colunas = 1 + len(dias_ordenados) + 4  # NOMES + dias + TOTAL + ENTREGAS + VALOR + A PAGAR
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=num_colunas)
        title_cell = ws['A1']
        title_cell.value = f"📊 ENTREGADORES - {periodo.upper()}"
        title_cell.font = title_font
        title_cell.fill = title_fill
        title_cell.alignment = title_align
        title_cell.border = thin_border
        ws.row_dimensions[1].height = 25
        
        # Linha 2: Cabeçalho "NOMES" + dias da semana
        col_num = 1
        
        # Coluna A: "NOMES"
        cell = ws.cell(row=2, column=col_num, value="NOMES")
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border
        ws.column_dimensions['A'].width = 20
        col_num += 1
        
        # Colunas dos dias da semana/datas
        for dia in dias_ordenados:
            if formato_data:
                # Para datas, usa a data diretamente
                dia_nome = dia
            else:
                # Para nomes de dias, usa o mapeamento
                dia_nome = dias_semana_map.get(dia, dia.capitalize())
            
            cell = ws.cell(row=2, column=col_num, value=dia_nome)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
            ws.column_dimensions[get_column_letter(col_num)].width = 12
            col_num += 1
        
        # Últimas colunas: TOTAL, ENTREGAS, VALOR, A PAGAR
        for header in ["TOTAL", "ENTREGAS", "VALOR", "A PAGAR"]:
            cell = ws.cell(row=2, column=col_num, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
            ws.column_dimensions[get_column_letter(col_num)].width = 12
            col_num += 1
        
        ws.row_dimensions[2].height = 20
        
        # Linhas 3+: Dados dos entregadores
        row_num = 3
        alt_fill = PatternFill(start_color="DAEEF3", end_color="DAEEF3", fill_type="solid")
        
        for idx, nome in enumerate(lista_entregadores):
            # Coluna A: Nome do entregador
            cell = ws.cell(row=row_num, column=1, value=nome)
            cell.alignment = Alignment(horizontal="left", vertical="center")
            cell.border = thin_border
            if idx % 2 == 1:
                cell.fill = alt_fill
            
            # Colunas dos dias: preenche com número de entregas
            col_num = 2
            for dia in dias_ordenados:
                # Busca entregas desse dia/pessoa
                entregas = 0
                if dia in entregas_por_dia_pessoa:
                    entregas = entregas_por_dia_pessoa[dia].get(nome, 0)
                
                cell = ws.cell(row=row_num, column=col_num, value=entregas if entregas > 0 else "")
                cell.alignment = data_align
                cell.border = thin_border
                if idx % 2 == 1:
                    cell.fill = alt_fill
                
                col_num += 1
            
            # Coluna TOTAL: soma dos dias
            primeira_col_dia = get_column_letter(2)
            ultima_col_dia = get_column_letter(1 + len(dias_ordenados))
            formula = f"=SUM({primeira_col_dia}{row_num}:{ultima_col_dia}{row_num})"
            cell = ws.cell(row=row_num, column=col_num, value=formula)
            cell.alignment = data_align
            cell.border = thin_border
            cell.font = Font(bold=True)
            if idx % 2 == 1:
                cell.fill = alt_fill
            col_num += 1
            
            # Coluna ENTREGAS: igual ao TOTAL
            col_total = get_column_letter(col_num - 1)
            formula = f"={col_total}{row_num}"
            cell = ws.cell(row=row_num, column=col_num, value=formula)
            cell.alignment = data_align
            cell.border = thin_border
            if idx % 2 == 1:
                cell.fill = alt_fill
            col_num += 1
            
            # Coluna VALOR: ENTREGAS * custo_entrega
            col_entregas = get_column_letter(col_num - 1)
            formula = f"={col_entregas}{row_num}*{custo_entrega}"
            cell = ws.cell(row=row_num, column=col_num, value=formula)
            cell.number_format = 'R$ #,##0.00'
            cell.alignment = data_align
            cell.border = thin_border
            if idx % 2 == 1:
                cell.fill = alt_fill
            col_num += 1
            
            # Coluna A PAGAR: deixa vazio (será preenchido manualmente)
            cell = ws.cell(row=row_num, column=col_num, value="")
            cell.alignment = data_align
            cell.border = thin_border
            if idx % 2 == 1:
                cell.fill = alt_fill
            
            row_num += 1
        
        # Linha TOTAL (soma de todos os entregadores)
        total_row = row_num
        total_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        total_font = Font(bold=True, size=11)
        
        # Coluna A: "TOTAL"
        cell = ws.cell(row=total_row, column=1, value="TOTAL")
        cell.font = total_font
        cell.fill = total_fill
        cell.alignment = data_align
        cell.border = thin_border
        
        # Colunas dos dias: soma
        col_num = 2
        for _ in dias_ordenados:
            col_letter = get_column_letter(col_num)
            formula = f"=SUM({col_letter}3:{col_letter}{total_row-1})"
            cell = ws.cell(row=total_row, column=col_num, value=formula)
            cell.font = total_font
            cell.fill = total_fill
            cell.alignment = data_align
            cell.border = thin_border
            col_num += 1
        
        # Colunas finais: somas (TOTAL, ENTREGAS, VALOR, A PAGAR)
        # CORREÇÃO #2: Aplicar formato moeda correto em cada coluna
        for i in range(4):  # TOTAL, ENTREGAS, VALOR, A PAGAR
            col_letter = get_column_letter(col_num)
            formula = f"=SUM({col_letter}3:{col_letter}{total_row-1})"
            cell = ws.cell(row=total_row, column=col_num, value=formula)

            # CORREÇÃO #2: Formato moeda para VALOR e A PAGAR (índices 2 e 3)
            if i >= 2:  # VALOR e A PAGAR são valores monetários
                cell.number_format = 'R$ #,##0.00'
                cell.alignment = currency_align
            else:
                cell.number_format = '#,##0'
                cell.alignment = data_align

            cell.font = total_font
            cell.fill = total_fill
            cell.border = thick_border
            col_num += 1
        
        ws.row_dimensions[total_row].height = 20
        
        # ===== ADICIONA GRÁFICOS =====
        try:
            from openpyxl.chart import BarChart, PieChart, Reference
            from openpyxl.chart.label import DataLabelList
            
            # Calcula posição inicial dos gráficos (após as colunas de dados)
            col_grafico_inicio = len(dias_ordenados) + 7  # Após NOMES + dias + TOTAL + ENTREGAS + VALOR + A PAGAR
            col_grafico_letra = get_column_letter(col_grafico_inicio)
            
            # GRÁFICO 1: Top 10 Entregadores (Barras Horizontais)
            chart1 = BarChart()
            chart1.type = "bar"
            chart1.style = 11
            chart1.title = "🏆 Top 10 Entregadores"
            chart1.y_axis.title = 'Entregador'
            chart1.x_axis.title = 'Entregas'

            # CORREÇÃO #8: Ordena entregadores por total de entregas (para gráfico Top 10)
            # Calcula totais de cada entregador
            totais_entregadores = []
            primeira_col_dia = 2
            ultima_col_dia = 1 + len(dias_ordenados)
            col_total_idx = 1 + len(dias_ordenados) + 1  # Coluna do TOTAL

            for idx, nome in enumerate(lista_entregadores):
                row = 3 + idx
                # Soma das entregas do entregador
                total = 0
                for col in range(primeira_col_dia, ultima_col_dia + 1):
                    cell_val = ws.cell(row=row, column=col).value
                    if cell_val and isinstance(cell_val, (int, float)):
                        total += cell_val
                totais_entregadores.append((nome, total, row))

            # Ordena por total (maior para menor)
            totais_entregadores.sort(key=lambda x: x[1], reverse=True)

            # Pega os top N para o gráfico
            num_entregadores_grafico = min(10, len(totais_entregadores))
            top_entregadores = totais_entregadores[:num_entregadores_grafico]

            # Prepara referências ordenadas para o gráfico
            #.openpyxl precisa de listas separadas de nomes e valores
            nomes_ordenados = [item[0] for item in top_entregadores]
            valores_ordenados = [item[1] for item in top_entregadores]
            linhas_ordenadas = [item[2] for item in top_entregadores]

            # Cria área temporária para o gráfico ordenado (colunas auxiliares após gráficos)
            col_temp = col_grafico_inicio + 15  # Área bem à direita
            ws.cell(row=2, column=col_temp, value="Ranking")
            ws.cell(row=2, column=col_temp + 1, value="Entregas")

            for i, (nome, total, _) in enumerate(top_entregadores):
                ws.cell(row=3 + i, column=col_temp, value=nome)
                ws.cell(row=3 + i, column=col_temp + 1, value=total)

            # Dados: Nomes e Total de Entregas (usando área temporária ordenada)
            col_temp_letter = get_column_letter(col_temp)
            col_temp_val = get_column_letter(col_temp + 1)
            data1 = Reference(ws, min_col=col_temp_val, min_row=2, max_row=2 + num_entregadores_grafico)
            cats1 = Reference(ws, min_col=col_temp_letter, min_row=3, max_row=2 + num_entregadores_grafico)
            chart1.add_data(data1, titles_from_data=True)
            chart1.set_categories(cats1)
            
            # Estilo
            chart1.height = 12
            chart1.width = 18
            
            # Posiciona o gráfico
            ws.add_chart(chart1, f"{col_grafico_letra}2")
            
            # GRÁFICO 2: Entregas por Dia (Colunas)
            chart2 = BarChart()
            chart2.type = "col"
            chart2.style = 10
            chart2.title = "📊 Entregas por Dia"
            chart2.y_axis.title = 'Quantidade'
            chart2.x_axis.title = 'Dia'
            
            # Dados: Dias (colunas 2 até len(dias_ordenados)+1) na linha TOTAL
            data2 = Reference(ws, min_col=2, min_row=total_row, max_col=len(dias_ordenados)+1, max_row=total_row)
            cats2 = Reference(ws, min_col=2, min_row=2, max_col=len(dias_ordenados)+1, max_row=2)
            chart2.add_data(data2, titles_from_data=True)
            chart2.set_categories(cats2)
            
            # Estilo
            chart2.height = 10
            chart2.width = 20
            
            # Posiciona o gráfico abaixo do primeiro
            ws.add_chart(chart2, f"{col_grafico_letra}20")
            
            # GRÁFICO 3: Distribuição de Entregas (Pizza) - Top 5
            if len(lista_entregadores) >= 3:
                chart3 = PieChart()
                chart3.title = "🍕 Distribuição de Entregas (Top 5)"
                
                # Pega os 5 primeiros entregadores
                num_top5 = min(5, len(lista_entregadores))
                
                # Dados: Total de entregas dos top 5
                data3 = Reference(ws, min_col=col_total, min_row=3, max_row=2+num_top5)
                labels3 = Reference(ws, min_col=1, min_row=3, max_row=2+num_top5)
                
                chart3.add_data(data3, titles_from_data=False)
                chart3.set_categories(labels3)
                
                # Mostra valores e porcentagens
                chart3.dataLabels = DataLabelList()
                chart3.dataLabels.showPercent = True
                chart3.dataLabels.showVal = True
                
                # Estilo
                chart3.height = 10
                chart3.width = 15
                
                # Posiciona o gráfico
                col_grafico2_letra = get_column_letter(col_grafico_inicio + 12)
                ws.add_chart(chart3, f"{col_grafico2_letra}20")
            
            logger.info("✅ Gráficos adicionados à planilha COM NOMES")
            
        except ImportError:
            logger.warning("⚠️ openpyxl.chart não disponível - gráficos não adicionados")
        except Exception as e:
            logger.warning(f"⚠️ Erro ao adicionar gráficos: {e}")
            import traceback
            logger.warning(traceback.format_exc())
        
        # Salva em bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"XLSX de entregadores com nomes criado: {len(result)} bytes, {len(lista_entregadores)} entregadores, {len(dias_ordenados)} dias")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar XLSX de entregadores com nomes: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None



# ============ FUNÇÃO PARA PLANILHAS PERSONALIZADAS ============

def criar_xlsx_estruturada(estrutura: dict, dados: list = None) -> Optional[bytes]:
    """Cria planilha Excel estruturada baseada em estrutura definida
    
    Args:
        estrutura: Dicionário com estrutura da planilha:
            {
                "titulo": "Nome da Planilha",
                "colunas": [
                    {"nome": "Col1", "tipo": "texto|numero|moeda|data", "largura": 15},
                    ...
                ],
                "tem_total": True/False,
                "colunas_total": ["Col1", "Col2"]
            }
        dados: Lista de listas com dados (opcional)
    
    Returns:
        bytes do XLSX ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        # Valida estrutura
        if not estrutura or not isinstance(estrutura, dict):
            logger.error("Estrutura inválida: deve ser um dicionário")
            return None
        
        if 'colunas' not in estrutura or not isinstance(estrutura['colunas'], list):
            logger.error("Estrutura inválida: falta 'colunas' ou não é lista")
            return None
        
        colunas = estrutura.get('colunas', [])
        num_colunas = len(colunas)
        
        if num_colunas == 0:
            logger.error("Estrutura inválida: nenhuma coluna definida")
            return None
        wb = Workbook()
        ws = wb.active
        ws.title = estrutura.get('titulo', 'Planilha')[:31]  # Excel limita a 31 caracteres
        
        # Estilos
        title_font = Font(bold=True, size=14, color="FFFFFF")
        title_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        title_align = Alignment(horizontal="center", vertical="center")
        
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        
        data_align = Alignment(horizontal="center", vertical="center")
        
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        thick_border = Border(
            left=Side(style='medium'),
            right=Side(style='medium'),
            top=Side(style='medium'),
            bottom=Side(style='medium')
        )
        
        # Linha 1: Título
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=num_colunas)
        title_cell = ws['A1']
        title_cell.value = f"📊 {estrutura.get('titulo', 'PLANILHA').upper()}"
        title_cell.font = title_font
        title_cell.fill = title_fill
        title_cell.alignment = title_align
        title_cell.border = thick_border
        ws.row_dimensions[1].height = 30
        
        # Linha 2: Vazia (espaçamento)
        ws.row_dimensions[2].height = 5
        
        # Linha 3: Cabeçalhos
        for col_num, coluna in enumerate(colunas, 1):
            cell = ws.cell(row=3, column=col_num, value=coluna['nome'])
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
            
            # Ajusta largura
            largura = coluna.get('largura', 15)
            ws.column_dimensions[get_column_letter(col_num)].width = largura
        
        ws.row_dimensions[3].height = 25
        
        # Dados
        alt_fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
        
        if dados:
            for row_idx, linha_dados in enumerate(dados, 4):
                for col_idx, (valor, coluna) in enumerate(zip(linha_dados, colunas), 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.border = thin_border
                    
                    # Converte e valida tipo de dado
                    tipo = coluna.get('tipo', 'texto')
                    try:
                        if tipo == 'moeda' or tipo == 'numero':
                            # Tenta converter para número
                            if isinstance(valor, str):
                                valor_limpo = valor.replace('R$', '').replace('.', '').replace(',', '.').strip()
                                valor = float(valor_limpo) if valor_limpo else 0
                            elif not isinstance(valor, (int, float)):
                                valor = 0
                            cell.value = valor
                            
                            if tipo == 'moeda':
                                cell.number_format = 'R$ #,##0.00'
                                cell.alignment = Alignment(horizontal="right", vertical="center")
                            else:
                                cell.number_format = '#,##0'
                                cell.alignment = data_align
                        
                        elif tipo == 'porcentagem':
                            # Tenta converter para número (0.15 = 15%)
                            if isinstance(valor, str):
                                valor_limpo = valor.replace('%', '').strip()
                                valor = float(valor_limpo) / 100 if valor_limpo else 0
                            elif not isinstance(valor, (int, float)):
                                valor = 0
                            cell.value = valor
                            cell.number_format = '0.00%'
                            cell.alignment = data_align
                        
                        elif tipo == 'data':
                            # Mantém como string (formato DD/MM/YYYY)
                            cell.value = str(valor) if valor else ''
                            cell.number_format = 'DD/MM/YYYY'
                            cell.alignment = data_align
                        
                        else:  # texto
                            cell.value = str(valor) if valor else ''
                            cell.alignment = Alignment(horizontal="left", vertical="center")
                    
                    except (ValueError, AttributeError) as e:
                        # Se falhar conversão, usa como texto
                        logger.warning(f"Erro ao converter valor '{valor}' para tipo '{tipo}': {e}")
                        cell.value = str(valor) if valor else ''
                        cell.alignment = Alignment(horizontal="left", vertical="center")
                    
                    # Zebra stripes
                    if (row_idx - 4) % 2 == 1:
                        cell.fill = alt_fill
        
        # Linha de TOTAL (se necessário)
        if estrutura.get('tem_total', False) and dados:
            total_row = len(dados) + 4
            total_fill = PatternFill(start_color="27AE60", end_color="27AE60", fill_type="solid")
            total_font = Font(bold=True, color="FFFFFF", size=12)
            
            colunas_total = estrutura.get('colunas_total', [])
            
            for col_idx, coluna in enumerate(colunas, 1):
                cell = ws.cell(row=total_row, column=col_idx)
                cell.font = total_font
                cell.fill = total_fill
                cell.border = thick_border
                
                if col_idx == 1:
                    cell.value = "TOTAL"
                    cell.alignment = data_align
                elif coluna['nome'] in colunas_total:
                    # Cria fórmula de soma
                    col_letter = get_column_letter(col_idx)
                    formula = f"=SUM({col_letter}4:{col_letter}{total_row-1})"
                    cell.value = formula
                    
                    # Formato baseado no tipo
                    if coluna.get('tipo') == 'moeda':
                        cell.number_format = 'R$ #,##0.00'
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                    else:
                        cell.number_format = '#,##0'
                        cell.alignment = data_align
                else:
                    cell.value = ""
                    cell.alignment = data_align
            
            ws.row_dimensions[total_row].height = 25
        
        # Salva em bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        
        result = buffer.getvalue()
        buffer.close()
        
        logger.info(f"XLSX estruturada criada: {len(result)} bytes, {num_colunas} colunas, {len(dados) if dados else 0} linhas")
        return result
        
    except Exception as e:
        logger.error(f"Erro ao criar XLSX estruturada: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None


def aplicar_edicao_planilha(xlsx_bytes: bytes, acao: str, parametros: dict, estrutura: dict) -> Optional[Tuple[bytes, str]]:
    """Aplica edição em planilha Excel existente
    
    Args:
        xlsx_bytes: bytes da planilha original
        acao: Tipo de ação (adicionar_linha|editar_celula|remover_linha|editar_coluna|substituir_valor)
        parametros: Parâmetros da ação
        estrutura: Estrutura da planilha (colunas, tipos)
    
    Returns:
        Tuple (bytes da planilha modificada, mensagem de sucesso) ou None se falhar
    """
    if not HAS_XLSX:
        logger.error("openpyxl não instalado")
        return None
    
    try:
        # Carrega workbook
        wb = load_workbook(io.BytesIO(xlsx_bytes))
        ws = wb.active
        
        # Estilos padrão para formatação
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        data_align = Alignment(horizontal="center", vertical="center")
        
        # Identifica linha de dados (pula título e cabeçalho)
        primeira_linha_dados = 4
        
        if acao == 'adicionar_linha':
            # Adiciona nova linha de dados
            valores = parametros.get('valores', [])
            if not valores:
                return None
            
            # Encontra próxima linha vazia
            proxima_linha = ws.max_row + 1
            
            # Se tem linha de TOTAL, insere uma nova linha ANTES dela
            if estrutura.get('tem_total', False):
                linha_total = ws.max_row
                ws.insert_rows(linha_total)  # Insere linha vazia antes do TOTAL
                proxima_linha = linha_total  # Nova linha de dados
            
            # Adiciona valores
            colunas = estrutura.get('colunas', [])
            for col_idx, (valor, coluna) in enumerate(zip(valores, colunas), 1):
                cell = ws.cell(row=proxima_linha, column=col_idx)
                
                # Converte tipo
                tipo = coluna.get('tipo', 'texto')
                try:
                    if tipo in ['moeda', 'numero']:
                        if isinstance(valor, str):
                            valor_limpo = valor.replace('R$', '').replace('.', '').replace(',', '.').strip()
                            cell.value = float(valor_limpo) if valor_limpo else 0
                        else:
                            cell.value = float(valor) if valor else 0
                        
                        if tipo == 'moeda':
                            cell.number_format = 'R$ #,##0.00'
                        else:
                            cell.number_format = '#,##0'
                    
                    elif tipo == 'porcentagem':
                        if isinstance(valor, str):
                            valor_limpo = valor.replace('%', '').strip()
                            cell.value = float(valor_limpo) / 100 if valor_limpo else 0
                        else:
                            cell.value = float(valor) / 100 if valor else 0
                        cell.number_format = '0.00%'
                    
                    else:
                        cell.value = str(valor) if valor else ''
                
                except (ValueError, AttributeError):
                    cell.value = str(valor) if valor else ''
                
                # Copia formatação da linha anterior (se existir)
                if proxima_linha > primeira_linha_dados:
                    cell_anterior = ws.cell(row=proxima_linha-1, column=col_idx)
                    cell.font = cell_anterior.font.copy()
                    cell.border = cell_anterior.border.copy()
                    cell.alignment = cell_anterior.alignment.copy()
                else:
                    # Primeira linha de dados - aplica formatação padrão
                    cell.border = thin_border
                    cell.alignment = data_align
                
                # Aplica zebra stripes (alternar cores)
                if estrutura.get('tem_total', False):
                    # Conta linhas de dados (excluindo título, cabeçalho e TOTAL)
                    num_dados = proxima_linha - primeira_linha_dados
                else:
                    num_dados = proxima_linha - primeira_linha_dados + 1
                
                if num_dados % 2 == 0:  # Linha par = zebra
                    cell.fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
            
            mensagem = f"✅ Linha adicionada com sucesso!"
        
        elif acao == 'editar_celula':
            # Edita célula específica
            linha = parametros.get('linha')
            coluna_nome = parametros.get('coluna')
            valores = parametros.get('valores', [])
            
            if not linha or not valores:
                return None
            
            # Converte linha para índice real (usuário conta a partir de 1, mas pula título/cabeçalho)
            linha_real = primeira_linha_dados + int(linha) - 1
            
            # Encontra índice da coluna
            colunas = estrutura.get('colunas', [])
            col_idx = None
            coluna_info = None
            
            for idx, col in enumerate(colunas, 1):
                if col['nome'].lower() == coluna_nome.lower():
                    col_idx = idx
                    coluna_info = col
                    break
            
            if not col_idx:
                # Se não encontrou por nome, tenta por índice
                try:
                    col_idx = int(coluna_nome)
                    coluna_info = colunas[col_idx - 1]
                except (ValueError, IndexError):
                    return None
            
            # Edita célula
            cell = ws.cell(row=linha_real, column=col_idx)
            novo_valor = valores[0]
            
            # Converte tipo
            tipo = coluna_info.get('tipo', 'texto')
            try:
                if tipo in ['moeda', 'numero']:
                    if isinstance(novo_valor, str):
                        valor_limpo = novo_valor.replace('R$', '').replace('.', '').replace(',', '.').strip()
                        cell.value = float(valor_limpo) if valor_limpo else 0
                    else:
                        cell.value = float(novo_valor) if novo_valor else 0
                elif tipo == 'porcentagem':
                    if isinstance(novo_valor, str):
                        valor_limpo = novo_valor.replace('%', '').strip()
                        cell.value = float(valor_limpo) / 100 if valor_limpo else 0
                    else:
                        cell.value = float(novo_valor) / 100 if novo_valor else 0
                else:
                    cell.value = str(novo_valor) if novo_valor else ''
            except (ValueError, AttributeError):
                cell.value = str(novo_valor) if novo_valor else ''
            
            mensagem = f"✅ Célula editada: linha {linha}, coluna {coluna_nome}"
        
        elif acao == 'remover_linha':
            # Remove linha
            linha = parametros.get('linha')
            
            if not linha:
                return None
            
            if linha == 'ultima':
                # Remove última linha de dados (antes do TOTAL se houver)
                linha_remover = ws.max_row
                if estrutura.get('tem_total', False):
                    linha_remover = ws.max_row - 1
            else:
                # Converte para índice real
                linha_remover = primeira_linha_dados + int(linha) - 1
            
            ws.delete_rows(linha_remover, 1)
            mensagem = f"✅ Linha removida com sucesso!"
        
        elif acao == 'editar_coluna':
            # Aplica operação em coluna inteira
            coluna_nome = parametros.get('coluna')
            operacao = parametros.get('operacao')
            fator = parametros.get('fator', 1)
            
            if not coluna_nome or not operacao:
                return None
            
            # Encontra índice da coluna
            colunas = estrutura.get('colunas', [])
            col_idx = None
            
            for idx, col in enumerate(colunas, 1):
                if col['nome'].lower() == coluna_nome.lower():
                    col_idx = idx
                    break
            
            if not col_idx:
                return None
            
            # Aplica operação em todas as linhas de dados
            ultima_linha = ws.max_row
            if estrutura.get('tem_total', False):
                ultima_linha -= 1  # Não modifica linha de TOTAL
            
            for linha in range(primeira_linha_dados, ultima_linha + 1):
                cell = ws.cell(row=linha, column=col_idx)
                if cell.value and isinstance(cell.value, (int, float)):
                    if operacao == 'multiplicar':
                        cell.value = cell.value * fator
                    elif operacao == 'dividir':
                        cell.value = cell.value / fator if fator != 0 else cell.value
                    elif operacao == 'somar':
                        cell.value = cell.value + fator
                    elif operacao == 'subtrair':
                        cell.value = cell.value - fator
            
            mensagem = f"✅ Coluna '{coluna_nome}' atualizada ({operacao} por {fator})"
        
        elif acao == 'substituir_valor':
            # Busca e substitui valores
            coluna_nome = parametros.get('coluna')
            valores = parametros.get('valores', [])
            
            if not valores or len(valores) < 2:
                return None
            
            valor_antigo = str(valores[0])
            valor_novo = str(valores[1])
            
            substituicoes = 0
            
            # Se especificou coluna, busca só nela
            if coluna_nome:
                colunas = estrutura.get('colunas', [])
                col_idx = None
                
                for idx, col in enumerate(colunas, 1):
                    if col['nome'].lower() == coluna_nome.lower():
                        col_idx = idx
                        break
                
                if col_idx:
                    ultima_linha = ws.max_row
                    if estrutura.get('tem_total', False):
                        ultima_linha -= 1
                    
                    for linha in range(primeira_linha_dados, ultima_linha + 1):
                        cell = ws.cell(row=linha, column=col_idx)
                        if cell.value and str(cell.value) == valor_antigo:
                            cell.value = valor_novo
                            substituicoes += 1
            else:
                # Busca em todas as colunas
                ultima_linha = ws.max_row
                if estrutura.get('tem_total', False):
                    ultima_linha -= 1
                
                for linha in range(primeira_linha_dados, ultima_linha + 1):
                    for col_idx in range(1, len(estrutura.get('colunas', [])) + 1):
                        cell = ws.cell(row=linha, column=col_idx)
                        if cell.value and str(cell.value) == valor_antigo:
                            cell.value = valor_novo
                            substituicoes += 1
            
            mensagem = f"✅ {substituicoes} substituição(ões) realizada(s)"
        
        else:
            logger.error(f"Ação desconhecida: {acao}")
            return None
        
        # Salva workbook modificado
        output = io.BytesIO()
        wb.save(output)
        wb.close()
        
        result = output.getvalue()
        output.close()
        
        logger.info(f"Edição aplicada: {acao}, {len(result)} bytes")
        return result, mensagem
        
    except Exception as e:
        logger.error(f"Erro ao aplicar edição: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None


# ============ FUNÇÕES AVANÇADAS DE DOCX (FASE 1) ============

def ler_docx_headers_footers(docx_bytes: bytes) -> Optional[dict]:
    """Lê headers e footers de um documento Word
    
    Args:
        docx_bytes: bytes do arquivo DOCX
    
    Returns:
        dict com headers e footers ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        resultado = {
            'headers': [],
            'footers': [],
            'num_secoes': len(doc.sections)
        }
        
        for i, section in enumerate(doc.sections):
            # Header
            header = section.header
            if header and not header.is_linked_to_previous:
                header_text = '\n'.join([p.text for p in header.paragraphs if p.text.strip()])
                if header_text:
                    resultado['headers'].append({
                        'secao': i,
                        'texto': header_text
                    })
            
            # Footer
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                footer_text = '\n'.join([p.text for p in footer.paragraphs if p.text.strip()])
                if footer_text:
                    resultado['footers'].append({
                        'secao': i,
                        'texto': footer_text
                    })
        
        logger.info(f"Headers/Footers lidos: {len(resultado['headers'])} headers, {len(resultado['footers'])} footers")
        return resultado
        
    except Exception as e:
        logger.error(f"Erro ao ler headers/footers: {e}")
        return None


def editar_docx_header(docx_bytes: bytes, novo_texto: str, secao_idx: int = 0) -> Optional[bytes]:
    """Edita o header de uma seção do documento
    
    Args:
        docx_bytes: bytes do arquivo DOCX original
        novo_texto: novo texto para o header
        secao_idx: índice da seção (padrão: 0 = primeira)
    
    Returns:
        bytes do DOCX modificado ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        if secao_idx >= len(doc.sections):
            logger.error(f"Seção {secao_idx} não existe (total: {len(doc.sections)})")
            return None
        
        section = doc.sections[secao_idx]
        header = section.header
        
        # Limpa header existente
        for p in header.paragraphs:
            p.clear()
        
        # Adiciona novo texto
        if header.paragraphs:
            header.paragraphs[0].add_run(novo_texto)
        else:
            header.add_paragraph(novo_texto)
        
        # Salva
        output = io.BytesIO()
        doc.save(output)
        
        logger.info(f"Header editado na seção {secao_idx}")
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao editar header: {e}")
        return None


def editar_docx_footer(docx_bytes: bytes, novo_texto: str, secao_idx: int = 0) -> Optional[bytes]:
    """Edita o footer de uma seção do documento
    
    Args:
        docx_bytes: bytes do arquivo DOCX original
        novo_texto: novo texto para o footer
        secao_idx: índice da seção (padrão: 0 = primeira)
    
    Returns:
        bytes do DOCX modificado ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        if secao_idx >= len(doc.sections):
            logger.error(f"Seção {secao_idx} não existe (total: {len(doc.sections)})")
            return None
        
        section = doc.sections[secao_idx]
        footer = section.footer
        
        # Limpa footer existente
        for p in footer.paragraphs:
            p.clear()
        
        # Adiciona novo texto
        if footer.paragraphs:
            footer.paragraphs[0].add_run(novo_texto)
        else:
            footer.add_paragraph(novo_texto)
        
        # Salva
        output = io.BytesIO()
        doc.save(output)
        
        logger.info(f"Footer editado na seção {secao_idx}")
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao editar footer: {e}")
        return None


def contar_imagens_docx(docx_bytes: bytes) -> int:
    """Conta imagens em um documento Word
    
    Args:
        docx_bytes: bytes do arquivo DOCX
    
    Returns:
        Número de imagens no documento
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return 0
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        count = 0
        
        # Conta imagens inline nos parágrafos
        for para in doc.paragraphs:
            for run in para.runs:
                if run._element.xpath('.//a:blip'):
                    count += 1
        
        # Conta imagens nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run._element.xpath('.//a:blip'):
                                count += 1
        
        logger.info(f"Imagens encontradas: {count}")
        return count
        
    except Exception as e:
        logger.error(f"Erro ao contar imagens: {e}")
        return 0


def validar_integridade_docx(docx_bytes: bytes, original_bytes: bytes = None) -> dict:
    """Valida integridade de um documento DOCX
    
    Args:
        docx_bytes: bytes do documento a validar
        original_bytes: bytes do documento original (opcional, para comparação)
    
    Returns:
        dict com resultado da validação
    """
    if not HAS_DOCX:
        return {'valido': False, 'erro': 'python-docx não instalado'}
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        resultado = {
            'valido': True,
            'num_paragrafos': len(doc.paragraphs),
            'num_tabelas': len(doc.tables),
            'num_secoes': len(doc.sections),
            'tem_imagens': False,
            'tem_headers': False,
            'tem_footers': False,
            'alertas': []
        }
        
        # Verifica imagens
        resultado['num_imagens'] = contar_imagens_docx(docx_bytes)
        resultado['tem_imagens'] = resultado['num_imagens'] > 0
        
        # Verifica headers/footers
        for section in doc.sections:
            if section.header and not section.header.is_linked_to_previous:
                resultado['tem_headers'] = True
            if section.footer and not section.footer.is_linked_to_previous:
                resultado['tem_footers'] = True
        
        # Compara com original se fornecido
        if original_bytes:
            doc_orig = Document(io.BytesIO(original_bytes))
            
            # Verifica se perdeu parágrafos
            if len(doc.paragraphs) < len(doc_orig.paragraphs) * 0.8:
                resultado['alertas'].append(f"Perda significativa de parágrafos: {len(doc_orig.paragraphs)} -> {len(doc.paragraphs)}")
            
            # Verifica se perdeu tabelas
            if len(doc.tables) < len(doc_orig.tables):
                resultado['alertas'].append(f"Perda de tabelas: {len(doc_orig.tables)} -> {len(doc.tables)}")
            
            # Verifica tamanho
            if len(docx_bytes) < len(original_bytes) * 0.3:
                resultado['alertas'].append(f"Redução drástica de tamanho: {len(original_bytes)} -> {len(docx_bytes)} bytes")
        
        # Verifica se tem alertas críticos
        if any('Perda' in a for a in resultado['alertas']):
            resultado['valido'] = False
        
        logger.info(f"Validação: {resultado}")
        return resultado
        
    except Exception as e:
        logger.error(f"Erro ao validar DOCX: {e}")
        return {'valido': False, 'erro': str(e)}


def editar_docx_preservar_imagens(docx_bytes: bytes, texto_antigo: str, texto_novo: str) -> Optional[Tuple[bytes, int]]:
    """Substitui texto preservando imagens existentes
    
    Versão melhorada de editar_docx_substituir que:
    1. Preserva imagens inline
    2. Preserva formatação de runs
    3. Valida integridade após edição
    
    Args:
        docx_bytes: bytes do arquivo DOCX original
        texto_antigo: texto a ser substituído
        texto_novo: novo texto
    
    Returns:
        Tuple (bytes do DOCX modificado, número de substituições) ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        substituicoes = 0
        
        # Substitui em parágrafos preservando runs
        for para in doc.paragraphs:
            if texto_antigo in para.text:
                # Conta substituições
                substituicoes += para.text.count(texto_antigo)
                
                # Substitui dentro de cada run (preserva formatação)
                for run in para.runs:
                    if texto_antigo in run.text:
                        run.text = run.text.replace(texto_antigo, texto_novo)
        
        # Substitui em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if texto_antigo in cell.text:
                        substituicoes += cell.text.count(texto_antigo)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if texto_antigo in run.text:
                                    run.text = run.text.replace(texto_antigo, texto_novo)
        
        # Salva
        output = io.BytesIO()
        doc.save(output)
        result = output.getvalue()
        
        # Valida integridade
        validacao = validar_integridade_docx(result, docx_bytes)
        if not validacao['valido']:
            logger.warning(f"Validação falhou: {validacao['alertas']}")
            # Retorna mesmo assim, mas loga o aviso
        
        logger.info(f"Substituição preservando imagens: {substituicoes} ocorrências")
        return result, substituicoes
        
    except Exception as e:
        logger.error(f"Erro ao substituir preservando imagens: {e}")
        return None


def adicionar_imagem_docx(docx_bytes: bytes, imagem_bytes: bytes, posicao: str = 'fim', 
                          largura_polegadas: float = 4.0) -> Optional[bytes]:
    """Adiciona imagem a um documento Word
    
    Args:
        docx_bytes: bytes do arquivo DOCX original
        imagem_bytes: bytes da imagem (PNG, JPG, etc)
        posicao: 'inicio' ou 'fim' (padrão: 'fim')
        largura_polegadas: largura da imagem em polegadas
    
    Returns:
        bytes do DOCX com imagem ou None se falhar
    """
    if not HAS_DOCX:
        logger.error("python-docx não instalado")
        return None
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        
        # Cria buffer para a imagem
        img_stream = io.BytesIO(imagem_bytes)
        
        if posicao == 'inicio':
            # Adiciona no início
            para = doc.paragraphs[0].insert_paragraph_before()
            run = para.add_run()
            run.add_picture(img_stream, width=Inches(largura_polegadas))
        else:
            # Adiciona no fim
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(img_stream, width=Inches(largura_polegadas))
        
        # Salva
        output = io.BytesIO()
        doc.save(output)
        
        logger.info(f"Imagem adicionada ao documento ({largura_polegadas} polegadas de largura)")
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao adicionar imagem: {e}")
        return None


def adicionar_imagem_flutuante(docx_bytes: bytes, imagem_bytes: bytes, 
                               pos_x: float, pos_y: float,
                               largura_polegadas: float = 2.0,
                               altura_polegadas: float = 2.0) -> Optional[bytes]:
    """Adiciona imagem flutuante em posição específica (requer skelmis)
    
    Args:
        docx_bytes: bytes do arquivo DOCX original
        imagem_bytes: bytes da imagem
        pos_x: posição X em pontos
        pos_y: posição Y em pontos
        largura_polegadas: largura em polegadas
        altura_polegadas: altura em polegadas
    
    Returns:
        bytes do DOCX com imagem flutuante ou None se falhar
    """
    if not HAS_SKELMIS_DOCX:
        logger.warning("Imagens flutuantes requerem skelmis-python-docx")
        # Fallback: adiciona imagem inline
        return adicionar_imagem_docx(docx_bytes, imagem_bytes, 'fim', largura_polegadas)
    
    try:
        from skelmis.docx.shared import Pt
        
        doc = Document(io.BytesIO(docx_bytes))
        img_stream = io.BytesIO(imagem_bytes)
        
        # Adiciona parágrafo com imagem flutuante
        para = doc.add_paragraph()
        run = para.add_run()
        
        # Usa add_float_picture do skelmis
        run.add_float_picture(
            img_stream,
            width=Inches(largura_polegadas),
            height=Inches(altura_polegadas),
            pos_x=Pt(pos_x),
            pos_y=Pt(pos_y)
        )
        
        # Salva
        output = io.BytesIO()
        doc.save(output)
        
        logger.info(f"Imagem flutuante adicionada em ({pos_x}, {pos_y})")
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao adicionar imagem flutuante: {e}")
        # Fallback para inline
        return adicionar_imagem_docx(docx_bytes, imagem_bytes, 'fim', largura_polegadas)
