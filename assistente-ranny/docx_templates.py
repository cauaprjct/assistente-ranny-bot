"""
📄 Sistema de Templates DOCX
Assistente Ranny V3

Funcionalidades:
- Templates Jinja2 para documentos Word
- Templates pré-definidos para uso comum
- Renderização com variáveis dinâmicas
"""

import io
import logging
import os
from datetime import datetime
from typing import Optional, Dict, Any

logger = logging.getLogger(__name__)

# Tenta importar docxtpl
try:
    from docxtpl import DocxTemplate
    HAS_DOCXTPL = True
    logger.info("docxtpl disponível para templates")
except ImportError:
    HAS_DOCXTPL = False
    logger.warning("docxtpl não instalado - templates desabilitados")

# Diretório de templates
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), 'templates')


# ============ TEMPLATES PRÉ-DEFINIDOS ============

TEMPLATES_DISPONIVEIS = {
    'contrato_entregador': {
        'descricao': 'Contrato de prestação de serviços para entregadores',
        'variaveis': ['nome_entregador', 'cpf', 'rg', 'data_inicio', 'valor_entrega', 'dias_trabalho'],
        'arquivo': 'contrato_entregador.docx'
    },
    'relatorio_semanal': {
        'descricao': 'Relatório semanal de entregas',
        'variaveis': ['periodo', 'total_entregas', 'total_entregadores', 'valor_total', 'dias'],
        'arquivo': 'relatorio_semanal.docx'
    },
    'comprovante_pagamento': {
        'descricao': 'Comprovante de pagamento para entregadores',
        'variaveis': ['nome', 'periodo', 'valor', 'data_pagamento', 'forma_pagamento'],
        'arquivo': 'comprovante_pagamento.docx'
    },
    'recibo_simples': {
        'descricao': 'Recibo simples de pagamento',
        'variaveis': ['valor', 'valor_extenso', 'referente', 'nome_recebedor', 'cpf_recebedor', 'data'],
        'arquivo': 'recibo_simples.docx'
    },
    'lista_presenca': {
        'descricao': 'Lista de presença para reuniões',
        'variaveis': ['titulo_reuniao', 'data_reuniao', 'participantes'],
        'arquivo': 'lista_presenca.docx'
    }
}


def listar_templates() -> Dict[str, Dict]:
    """Lista todos os templates disponíveis
    
    Returns:
        Dict com templates e suas informações
    """
    return TEMPLATES_DISPONIVEIS


def obter_template(nome_template: str) -> Optional[Dict]:
    """Obtém informações de um template específico
    
    Args:
        nome_template: Nome do template
    
    Returns:
        Dict com informações do template ou None se não existir
    """
    return TEMPLATES_DISPONIVEIS.get(nome_template)


def renderizar_template(nome_template: str, contexto: Dict[str, Any]) -> Optional[bytes]:
    """Renderiza um template com variáveis

    Args:
        nome_template: Nome do template (ex: 'contrato_entregador')
        contexto: Dict com variáveis para o template

    Returns:
        bytes do DOCX renderizado ou None se falhar
    """
    if not HAS_DOCXTPL:
        logger.error("docxtpl não instalado")
        return None

    template_info = TEMPLATES_DISPONIVEIS.get(nome_template)
    if not template_info:
        logger.error(f"Template não encontrado: {nome_template}")
        return None

    # CORREÇÃO #5: Valida variáveis obrigatórias
    variaveis_obrigatorias = template_info.get('variaveis', [])
    variaveis_faltantes = []

    for var in variaveis_obrigatorias:
        if var not in contexto or not contexto[var]:
            variaveis_faltantes.append(var)

    if variaveis_faltantes:
        logger.warning(f"Variáveis faltantes para template '{nome_template}': {variaveis_faltantes}")
        # Preenche com valores padrão para não quebrar
        for var in variaveis_faltantes:
            contexto[var] = f"[{var.upper()}]"  # Placeholder visível

    template_path = os.path.join(TEMPLATES_DIR, template_info['arquivo'])
    
    # Verifica se arquivo existe
    if not os.path.exists(template_path):
        logger.error(f"Arquivo de template não encontrado: {template_path}")
        # Tenta criar template básico
        return criar_template_basico(nome_template, contexto)
    
    try:
        # Carrega template
        doc = DocxTemplate(template_path)
        
        # Adiciona variáveis padrão se não fornecidas
        contexto_completo = {
            'data_hoje': datetime.now().strftime('%d/%m/%Y'),
            'hora_agora': datetime.now().strftime('%H:%M'),
            'ano_atual': datetime.now().year,
            **contexto  # Sobrescreve com variáveis do usuário
        }
        
        # Renderiza
        doc.render(contexto_completo)
        
        # Salva em bytes
        output = io.BytesIO()
        doc.save(output)
        
        logger.info(f"Template '{nome_template}' renderizado com sucesso")
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao renderizar template: {e}")
        return None


def criar_template_basico(nome_template: str, contexto: Dict[str, Any]) -> Optional[bytes]:
    """Cria um documento básico quando o template não existe
    
    Args:
        nome_template: Nome do template
        contexto: Variáveis para o documento
    
    Returns:
        bytes do DOCX ou None se falhar
    """
    # Esta função usa apenas python-docx (skelmis), não precisa de docxtpl
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Título baseado no template
        titulos = {
            'contrato_entregador': 'CONTRATO DE PRESTAÇÃO DE SERVIÇOS',
            'relatorio_semanal': 'RELATÓRIO SEMANAL DE ENTREGAS',
            'comprovante_pagamento': 'COMPROVANTE DE PAGAMENTO',
            'recibo_simples': 'RECIBO',
            'lista_presenca': 'LISTA DE PRESENÇA'
        }
        
        titulo = titulos.get(nome_template, 'DOCUMENTO')
        heading = doc.add_heading(titulo, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Adiciona conteúdo baseado no contexto
        if nome_template == 'contrato_entregador':
            criar_contrato_basico(doc, contexto)
        elif nome_template == 'relatorio_semanal':
            criar_relatorio_basico(doc, contexto)
        elif nome_template == 'comprovante_pagamento':
            criar_comprovante_basico(doc, contexto)
        elif nome_template == 'recibo_simples':
            criar_recibo_basico(doc, contexto)
        else:
            # Documento genérico
            for chave, valor in contexto.items():
                p = doc.add_paragraph()
                p.add_run(f"{chave.replace('_', ' ').title()}: ").bold = True
                p.add_run(str(valor))
        
        # Rodapé
        doc.add_paragraph()
        rodape = doc.add_paragraph()
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rodape.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')} • Assistente Ranny")
        run.font.size = Pt(8)
        run.font.italic = True
        
        # Salva
        output = io.BytesIO()
        doc.save(output)
        
        logger.info(f"Documento básico criado para template '{nome_template}'")
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao criar documento básico: {e}")
        return None


def criar_contrato_basico(doc, contexto: Dict):
    """Cria contrato básico para entregador"""
    doc.add_paragraph("Pelo presente instrumento particular, de um lado:")
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("CONTRATANTE: ").bold = True
    p.add_run("GRN PIZZAS, estabelecimento comercial situado à [endereço].")
    
    doc.add_paragraph("E de outro lado:")
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("CONTRATADO: ").bold = True
    p.add_run(f"{contexto.get('nome_entregador', '[NOME]')}, ")
    p.add_run(f"CPF: {contexto.get('cpf', '[CPF]')}, ")
    p.add_run(f"RG: {contexto.get('rg', '[RG]')}.")
    
    doc.add_paragraph()
    doc.add_paragraph("Celebram o presente contrato de prestação de serviços, mediante as seguintes cláusulas:")
    doc.add_paragraph()
    
    doc.add_paragraph("CLÁUSULA PRIMEIRA - DO OBJETO")
    doc.add_paragraph("O presente contrato tem por objeto a prestação de serviços de entrega de pizzas e demais produtos comercializados pelo CONTRATANTE.")
    
    doc.add_paragraph("CLÁUSULA SEGUNDA - DA REMUNERAÇÃO")
    p = doc.add_paragraph()
    p.add_run(f"Pela prestação dos serviços, o CONTRATADO receberá o valor de R$ {contexto.get('valor_entrega', '12,00')} por entrega realizada.")
    
    doc.add_paragraph("CLÁUSULA TERCEIRA - DO PERÍODO")
    p = doc.add_paragraph()
    p.add_run(f"O contrato terá início em {contexto.get('data_inicio', datetime.now().strftime('%d/%m/%Y'))} e vigorará por tempo indeterminado.")
    
    doc.add_paragraph("CLÁUSULA QUARTA - DAS OBRIGAÇÕES")
    doc.add_paragraph("O CONTRATADO se obriga a:")
    doc.add_paragraph("• Realizar as entregas com pontualidade e cuidado;", style='List Bullet')
    doc.add_paragraph("• Zelar pela integridade dos produtos durante o transporte;", style='List Bullet')
    doc.add_paragraph("• Apresentar-se devidamente uniformizado;", style='List Bullet')
    doc.add_paragraph("• Manter seu veículo em boas condições de uso.", style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph("E por estarem assim justos e contratados, firmam o presente instrumento em 02 (duas) vias de igual teor e forma.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Assinaturas
    p = doc.add_paragraph()
    p.add_run("_________________________").bold = True
    p.add_run("          ")
    p.add_run("_________________________").bold = True
    
    p = doc.add_paragraph()
    p.add_run("CONTRATANTE")
    p.add_run("                              ")
    p.add_run("CONTRATADO")


def criar_relatorio_basico(doc, contexto: Dict):
    """Cria relatório semanal básico"""
    p = doc.add_paragraph()
    p.add_run("Período: ").bold = True
    p.add_run(contexto.get('periodo', 'Não informado'))
    
    doc.add_paragraph()
    
    # Resumo
    doc.add_heading('Resumo', level=2)
    
    p = doc.add_paragraph()
    p.add_run("Total de Entregas: ").bold = True
    p.add_run(str(contexto.get('total_entregas', 0)))
    
    p = doc.add_paragraph()
    p.add_run("Total de Entregadores: ").bold = True
    p.add_run(str(contexto.get('total_entregadores', 0)))
    
    p = doc.add_paragraph()
    p.add_run("Valor Total: ").bold = True
    p.add_run(f"R$ {contexto.get('valor_total', '0,00')}")
    
    # Detalhamento por dia
    dias = contexto.get('dias', [])
    if dias:
        doc.add_heading('Detalhamento por Dia', level=2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Cabeçalho
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Dia'
        header_cells[1].text = 'Entregadores'
        header_cells[2].text = 'Entregas'
        header_cells[3].text = 'Valor'
        
        for dia in dias:
            row_cells = table.add_row().cells
            row_cells[0].text = str(dia.get('dia', ''))
            row_cells[1].text = str(dia.get('entregadores', 0))
            row_cells[2].text = str(dia.get('entregas', 0))
            row_cells[3].text = f"R$ {dia.get('valor', '0,00')}"


def criar_comprovante_basico(doc, contexto: Dict):
    """Cria comprovante de pagamento básico"""
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Nome: ").bold = True
    p.add_run(contexto.get('nome', '[NOME]'))
    
    p = doc.add_paragraph()
    p.add_run("Período: ").bold = True
    p.add_run(contexto.get('periodo', '[PERÍODO]'))
    
    p = doc.add_paragraph()
    p.add_run("Valor: ").bold = True
    p.add_run(f"R$ {contexto.get('valor', '0,00')}")
    
    p = doc.add_paragraph()
    p.add_run("Data do Pagamento: ").bold = True
    p.add_run(contexto.get('data_pagamento', datetime.now().strftime('%d/%m/%Y')))
    
    p = doc.add_paragraph()
    p.add_run("Forma de Pagamento: ").bold = True
    p.add_run(contexto.get('forma_pagamento', 'Não informado'))
    
    doc.add_paragraph()
    doc.add_paragraph("Declaro ter recebido o valor acima referente aos serviços prestados.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("_________________________")
    
    p = doc.add_paragraph()
    p.add_run("Assinatura")


def criar_recibo_basico(doc, contexto: Dict):
    """Cria recibo simples básico"""
    valor = contexto.get('valor', '0,00')
    valor_extenso = contexto.get('valor_extenso', '(valor por extenso)')
    
    p = doc.add_paragraph()
    p.add_run(f"R$ {valor}").bold = True
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run(f"Recebi de GRN PIZZAS a quantia de {valor_extenso}, ")
    p.add_run(f"referente a {contexto.get('referente', '[referente]')}.")
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Nome: ").bold = True
    p.add_run(contexto.get('nome_recebedor', '[NOME]'))
    
    p = doc.add_paragraph()
    p.add_run("CPF: ").bold = True
    p.add_run(contexto.get('cpf_recebedor', '[CPF]'))
    
    p = doc.add_paragraph()
    p.add_run("Data: ").bold = True
    p.add_run(contexto.get('data', datetime.now().strftime('%d/%m/%Y')))
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("_________________________")
    
    p = doc.add_paragraph()
    p.add_run("Assinatura")


def criar_template_personalizado(nome: str, descricao: str, variaveis: list, template_bytes: bytes) -> bool:
    """Salva um template personalizado
    
    Args:
        nome: Nome do template
        descricao: Descrição do template
        variaveis: Lista de variáveis necessárias
        template_bytes: Bytes do arquivo DOCX template
    
    Returns:
        True se salvo com sucesso, False caso contrário
    """
    try:
        # Cria diretório se não existir
        os.makedirs(TEMPLATES_DIR, exist_ok=True)
        
        # Salva arquivo
        template_path = os.path.join(TEMPLATES_DIR, f"{nome}.docx")
        with open(template_path, 'wb') as f:
            f.write(template_bytes)
        
        # Adiciona à lista de templates
        TEMPLATES_DISPONIVEIS[nome] = {
            'descricao': descricao,
            'variaveis': variaveis,
            'arquivo': f"{nome}.docx"
        }
        
        logger.info(f"Template '{nome}' salvo com sucesso")
        return True
        
    except Exception as e:
        logger.error(f"Erro ao salvar template: {e}")
        return False
